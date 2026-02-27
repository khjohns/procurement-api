"""Word document generator for Del III (EØS) anskaffelsesprotokoll."""

from __future__ import annotations

import re
from datetime import datetime

from docx import Document as DocxDocument

from .common import (
    ALL_PROCEDURES,
    PROCEDURE_MAP,
    build_org_lookup,
    fmt_currency,
    fmt_date,
    fmt_datetime,
    get_activities_by_action,
    get_org_name,
    get_timeline_date,
    parse_submission_deadline,
    strip_html,
)
from .docx_helpers import (
    add_checkbox,
    add_instruction,
    add_manual,
    docx_add_table,
    docx_add_table_with_manual,
    docx_info_table,
    docx_setup,
    docx_subtitle,
)


def _general_info(doc, procurement, activities, org_lookup):
    doc.add_heading("Generell informasjon", level=3)

    seq_id = procurement.get("sequenceId") or ""
    ext_id = procurement.get("externalId") or ""
    sak_ref = seq_id
    if ext_id:
        sak_ref += f" (ekstern ref: {ext_id})"

    procurer = procurement.get("about_procurer") or {}
    procurer_name = procurer.get("name") or ""
    national_id = procurer.get("national_id") or ""
    procurer_str = procurer_name
    if national_id:
        procurer_str += f" (org.nr. {national_id})"

    contact = procurer.get("contact_person") or ""

    name = strip_html(procurement.get("name") or "")
    description = strip_html(procurement.get("description") or "")
    description = re.sub(r"\s*\n\s*", " ", description)
    desc_str = name
    if description and description.lower() != name.lower():
        desc_str += f". {description}"

    estimated_value = procurement.get("estimated_value")
    currency = procurement.get("currency") or "NOK"
    value_str = fmt_currency(estimated_value, currency)

    duration_months = procurement.get("duration_months")
    duration = procurement.get("duration")
    if duration_months:
        duration_str = f"{duration_months} måneder"
    elif duration:
        duration_str = str(duration)
    else:
        duration_str = None

    submission_date = get_timeline_date(procurement, "submission")
    submission_str = fmt_datetime(submission_date) if submission_date else None

    docx_info_table(doc, [
        ("Konkurransens saksnummer:", sak_ref),
        ("Oppdragsgiver(e):", procurer_str),
        ("Protokollfører/saksbehandler:", contact or None),
        ("Beskrivelse av anskaffelsen:", desc_str),
        ("Anskaffelsens anslåtte verdi på kunngjøringstidspunktet:", value_str),
        ("Kontraktens varighet:", duration_str),
        ("Frist for innlevering av tilbud:", submission_str),
    ])

    submissions = get_activities_by_action(activities, "SUBMIT_BID")
    p = doc.add_paragraph()
    p.add_run("Tidspunkt for mottak av tilbud:").bold = True

    if submissions:
        rows = []
        for s in submissions:
            org_name = get_org_name(s, org_lookup)
            date = fmt_datetime(s.get("date"))
            rows.append((org_name, date))
        docx_add_table(doc, ["Leverandørens navn", "Tidspunkt for mottak"], rows)
    else:
        doc.add_paragraph("Ingen tilbud registrert.")


def _procedure(doc, procurement, activities):
    doc.add_heading("Anskaffelsesprosedyre", level=3)
    doc.add_paragraph("Følgende anskaffelsesprosedyre er lagt til grunn i denne konkurransen:")

    procedure = procurement.get("procedure") or ""
    selected = PROCEDURE_MAP.get(procedure, "")
    for p_name in ALL_PROCEDURES:
        p = doc.add_paragraph()
        add_checkbox(p, p_name, checked=(p_name == selected), bold_if_checked=True)

    # -- Begrunnelse for prosedyrevalg --
    if procedure in ("Competitive negotiated", "Competitive dialogue"):
        dialog_str = None  # triggers [Fyll inn]
    else:
        dialog_str = "Ikke relevant."

    if procedure in ("Negotiated without publication", "Direct award"):
        code = procurement.get("direct_award_justification_code") or ""
        reason = procurement.get("direct_award_justification_reason") or ""
        if code or reason:
            no_pub_str = f"Hjemmel: {code}. {reason}" if code else reason
        else:
            no_pub_str = None
    else:
        no_pub_str = "Ikke relevant."

    # -- Kunngjøring --
    doffin_activities = get_activities_by_action(activities, "DOFFIN_NOTICE_STATUS_PUBLISHED")
    publish_activities = get_activities_by_action(activities, "PUBLISH_TO_DOFFIN")
    announcement_date = ""
    doffin_ref = ""
    ted_ref = ""
    if publish_activities:
        announcement_date = fmt_date(publish_activities[0].get("date"))
    if doffin_activities:
        desc = doffin_activities[0].get("description") or {}
        doffin_notice = desc.get("doffinNotice") or {}
        doffin_ref = doffin_notice.get("ngoj") or ""
        ted_ref = doffin_notice.get("publicationId") or ""
        if not announcement_date:
            announcement_date = fmt_date(
                doffin_notice.get("publicationDate") or doffin_activities[0].get("date")
            )

    if announcement_date:
        parts = [f"Kunngjort {announcement_date} på Doffin"]
        if doffin_ref:
            parts.append(f"ref. {doffin_ref} (NGOJ)")
        if ted_ref:
            parts.append(f"TED {ted_ref}")
        kunngj_str = ", ".join(parts)
    else:
        pub_date = procurement.get("publicationDate")
        kunngj_str = f"Kunngjort {fmt_date(pub_date)}" if pub_date else None

    docx_info_table(doc, [
        ("Begrunnelse for dialog/forhandling med kunngjøring:", dialog_str),
        ("Begrunnelse for forhandling uten kunngjøring / uten konkurranse:", no_pub_str),
        ("Begrunnelse for ikke å dele opp i delkontrakter (FOA \u00a7 19-4):", None),
        ("Kunngjøring:", kunngj_str),
    ])

    # Betingede lovkrav (§ 22-4 og § 25-5 (2) bokstav f)
    p = doc.add_paragraph()
    add_checkbox(p, "Unntak fra krav om elektronisk kommunikasjon ved mottak av tilbud")
    add_manual(p, " [Fyll inn begrunnelse, jf. FOA \u00a7 22-4]")


def _formal_rejection(doc, activities, org_lookup):
    doc.add_heading("Avvisning på grunn av formalfeil, jf. FOA \u00a7 24-1", level=3)

    rejections = get_activities_by_action(activities, "REJECT_PARTICIPATION")
    if not rejections:
        p = doc.add_paragraph()
        add_checkbox(p, "Ingen leverandører eller tilbud ble avvist")
        add_manual(p, " [Bekreft]")
        return

    p = doc.add_paragraph()
    add_manual(p, "[Avgjør hvilke avvisninger som gjelder formalfeil (\u00a7 24-1) vs. kvalifikasjon (\u00a7 24-2)]")

    rows = []
    for r in rejections:
        org_name = get_org_name(r, org_lookup)
        date = fmt_date(r.get("date"))
        rows.append((org_name, None, date))
    docx_add_table_with_manual(doc,
        ["Leverandørens navn", "Begrunnelsen for avvisningen", "Dato sendt"], rows)


def _preliminary_qualification(doc, procedure):
    doc.add_heading("Foreløpig kvalifikasjonsvurdering, jf. FOA \u00a7 17-1 annet ledd", level=3)

    if procedure == "Open":
        doc.add_paragraph("Ikke relevant (åpen anbudskonkurranse).")
    else:
        p = doc.add_paragraph()
        p.add_run("Leverandører som er kvalifisert:").bold = True
        p2 = doc.add_paragraph()
        add_manual(p2)


def _qualification(doc):
    doc.add_heading("Kvalifikasjonsvurdering", level=3)
    p = doc.add_paragraph()
    add_manual(p, "[Kvalifikasjonskrav og -vurdering er ikke tilgjengelig via API. Fyll inn basert på konkurransegrunnlaget.]")

    docx_info_table(doc, [
        ("Leverandører som er kvalifisert:", None),
        ("Hvis relevant, begrunnelse for hvorfor leverandører som har restanser i henhold til skatte- og avgiftslovgivningen har fått delta i konkurransen:", None),
    ])


def _supplier_rejection(doc, activities, org_lookup):
    doc.add_heading("Leverandører som er avvist, jf. FOA \u00a7 24-2", level=3)

    rejections = get_activities_by_action(activities, "REJECT_PARTICIPATION")
    if not rejections:
        p = doc.add_paragraph()
        add_checkbox(p, "Ingen leverandører ble avvist")
        add_manual(p, " [Bekreft]")
        return

    doc.add_paragraph("Følgende leverandører ble avvist:")
    rows = []
    for r in rejections:
        org_name = get_org_name(r, org_lookup)
        date = fmt_date(r.get("date"))
        rows.append((org_name, None, date))
    docx_add_table_with_manual(doc,
        ["Leverandørens navn", "Begrunnelsen for avvisningen", "Dato sendt"], rows)


def _supplier_selection(doc, procedure, activities):
    doc.add_heading("Utvelgelse av leverandører", level=3)

    if procedure == "Open":
        doc.add_paragraph("Ikke relevant (åpen anbudskonkurranse — ingen utvelgelsesfase).")
    elif procedure in ("Limited", "Competitive negotiated", "Innovation partnership", "Competitive dialogue"):
        qualifying = get_activities_by_action(activities, "QUALIFYING_PARTICIPANTS")
        p = doc.add_paragraph()
        add_manual(p, "[Fyll inn begrunnelse for utvelgelse per leverandør]")
        if qualifying:
            rows = []
            for q in qualifying:
                desc = q.get("description") or {}
                tenders_ids = desc.get("tendersIds") or []
                for tid in tenders_ids:
                    rows.append((f"Leverandør (tender {tid})", None))
            if rows:
                docx_add_table_with_manual(doc,
                    ["Leverandørens navn", "Begrunnelse for utvelgelse"], rows)
        else:
            docx_add_table_with_manual(doc,
                ["Leverandørens navn", "Begrunnelse for utvelgelse"], [(None, None)])
    else:
        doc.add_paragraph("Ikke relevant.")


def _bid_rejection(doc):
    doc.add_heading("Tilbud som er avvist, jf. FOA \u00a7\u00a7 24-8 og 24-9", level=3)
    p = doc.add_paragraph()
    add_checkbox(p, "Ingen tilbud ble avvist")
    add_manual(p, " [Bekreft]")

    docx_add_table_with_manual(doc,
        ["Leverandørens navn", "Begrunnelsen for avvisningen", "Dato sendt"],
        [(None, None, None)])

    # Unormalt lave tilbud (§ 24-9)
    doc.add_heading("Unormalt lave tilbud, jf. FOA \u00a7 24-9", level=4)
    p = doc.add_paragraph()
    add_checkbox(p, "Ingen tilbud vurdert som unormalt lave")
    add_manual(p, " [Bekreft]")

    docx_info_table(doc, [
        ("Leverandør(er) bedt om redegjørelse:", None),
        ("Vurdering av redegjørelsen:", None),
    ])


def _clarification(doc, procurement, activities, org_lookup):
    doc.add_heading("Ettersending og avklaring av opplysninger og dokumentasjon, jf. FOA \u00a7 23-5", level=3)

    submission_deadline = parse_submission_deadline(procurement)
    conversations = get_activities_by_action(activities, "CONVERSATION_MARKED_COMPLETED")

    post_deadline_convs = []
    if submission_deadline and conversations:
        for c in conversations:
            conv_date_str = c.get("date")
            if conv_date_str:
                try:
                    conv_date = datetime.fromisoformat(conv_date_str.replace("Z", "+00:00"))
                    if conv_date > submission_deadline:
                        post_deadline_convs.append(c)
                except (ValueError, TypeError):
                    pass

    if not post_deadline_convs:
        p = doc.add_paragraph()
        add_checkbox(p, "Det ble ikke foretatt avklaringer eller dialog")
        if not conversations:
            add_manual(p, " [Bekreft \u2014 ingen avklaringshendelser i API]")
        else:
            add_manual(p, " [Bekreft \u2014 meldinger finnes, men alle er f\u00f8r tilbudsfrist (Q&A)]")
        return

    doc.add_paragraph("F\u00f8lgende avklaringer/ettersendinger ble gjennomf\u00f8rt etter tilbudsfrist:")
    rows = []
    for c in post_deadline_convs:
        org_name = get_org_name(c, org_lookup)
        date = fmt_date(c.get("date"))
        desc = c.get("description") or {}
        title = desc.get("conversationTitle") or ""
        how = f"Melding i KGV: \u00ab{title}\u00bb" if title else "Melding i KGV"
        rows.append((org_name, date, how))
    docx_add_table(doc, ["Leverandørens navn", "Dato", "Hvordan?"], rows)


def _negotiations(doc, procedure):
    doc.add_heading("Forhandlinger", level=3)

    if procedure in ("Competitive negotiated", "Innovation partnership"):
        p = doc.add_paragraph()
        add_manual(p, "[Fyll inn forhandlingsdetaljer]")
        p2 = doc.add_paragraph()
        add_checkbox(p2, "Det ble ikke gjennomført forhandlinger")
        docx_add_table_with_manual(doc,
            ["Leverandørens navn", "Dato for forhandling", "Mottatt revidert tilbud"],
            [(None, None, None)])
    else:
        label = PROCEDURE_MAP.get(procedure, procedure) if procedure != "Open" else "åpen anbudskonkurranse"
        doc.add_paragraph(f"Ikke relevant ({label}).")


def _dialog(doc, procedure):
    doc.add_heading("Dialog", level=3)

    if procedure == "Competitive dialogue":
        p = doc.add_paragraph()
        add_manual(p, "[Fyll inn dialogdetaljer]")
        p2 = doc.add_paragraph()
        add_checkbox(p2, "Det ble ikke gjennomført dialog")
        docx_add_table_with_manual(doc,
            ["Leverandørens navn", "Dato for dialog"],
            [(None, None)])
    else:
        label = PROCEDURE_MAP.get(procedure, procedure) if procedure != "Open" else "åpen anbudskonkurranse"
        doc.add_paragraph(f"Ikke relevant ({label}).")


def _bids_in_evaluation(doc, activities, org_lookup):
    doc.add_heading("Tilbud som er med i tildelingsvurderingen", level=3)

    submissions = get_activities_by_action(activities, "SUBMIT_BID")
    rejections = get_activities_by_action(activities, "REJECT_PARTICIPATION")
    withdrawals = get_activities_by_action(activities, "WITHDRAW_PARTICIPATION")

    rejected_names = {
        get_org_name(r, org_lookup).lower() for r in rejections
    }
    withdrawn_names = {
        get_org_name(w, org_lookup).lower() for w in withdrawals
    }
    excluded = rejected_names | withdrawn_names

    evaluated = []
    for s in submissions:
        name = get_org_name(s, org_lookup)
        if name.lower() not in excluded:
            evaluated.append(name)

    if evaluated:
        rows = [(str(i), name) for i, name in enumerate(evaluated, 1)]
    else:
        rows = [("", None)]
    docx_add_table_with_manual(doc, ["Tilbuds-/løpenummer", "Leverandørenes navn"], rows)


def _award(doc, procurement, activities):
    doc.add_heading("Det (de) valgte tilbud med begrunnelse og kontraktsverdi", level=3)
    p = doc.add_paragraph()
    add_manual(p, "[Fyll inn navn på valgt leverandør, tildelingsbegrunnelse og kontraktsverdi. API gir kun tenderIds, ikke leverandørnavn eller begrunnelse.]")

    total_value = procurement.get("contracts_total_value_amount")
    estimated = procurement.get("estimated_value")
    currency = procurement.get("currency") or "NOK"
    if total_value:
        value_str = fmt_currency(total_value, currency)
    elif estimated:
        value_str = f"{fmt_currency(estimated, currency)} (estimert verdi)"
    else:
        value_str = None

    award_letters = procurement.get("areAwardLettersSent")

    award_date = get_timeline_date(procurement, "award decision")
    if not award_date:
        award_activities = get_activities_by_action(activities, "AWARDING_PARTICIPANTS")
        if award_activities:
            award_date = award_activities[0].get("date")
    award_str = fmt_date(award_date) if award_date else None

    docx_info_table(doc, [
        ("Kontraktsverdi:", value_str),
        ("Tildelingsbeslutning:", award_str),
        ("Meddelelsesbrev sendt:", "Sendt (dato ikke tilgjengelig i API)" if award_letters else None),
        ("Karensperiodens utløp:", None),
        ("Eventuelle klager:", None),
        ("Resultat av klage:", None),
    ])


def _framework_agreement(doc, procurement):
    doc.add_heading("Tildeling av rammeavtaler", level=3)

    is_framework = procurement.get("framework_agreement_involved")
    if not is_framework:
        doc.add_paragraph("Ikke relevant (dette er ikke en rammeavtale).")
        return

    max_participants = procurement.get("framework_agreement_maximum_participants")
    if max_participants and int(max_participants) == 1:
        docx_info_table(doc, [
            ("Rammeavtale med én leverandør:", "Ja"),
            ("Rammeavtale med flere leverandører:", "Nei"),
        ])
    elif max_participants and int(max_participants) > 1:
        docx_info_table(doc, [
            ("Rammeavtale med én leverandør:", "Nei"),
            ("Rammeavtale med flere leverandører:", f"Ja (maks {max_participants} deltakere)"),
            ("Fordelingsmekanisme:", None),
            ("Ved minikonkurranse; hvilke kriterier:", None),
        ])
    else:
        docx_info_table(doc, [
            ("Rammeavtale med én leverandør:", None),
            ("Rammeavtale med flere leverandører:", None),
        ])


def _market_dialogue_and_conflicts(doc):
    """Dialog med markedet og habilitet, jf. FOA kap. 12 og § 7-5."""
    doc.add_heading("Dialog med markedet og habilitet", level=3)

    # Kap. 12 — Dialog med markedet før konkurranse
    doc.add_heading("Dialog med markedet f\u00f8r konkurranse, jf. FOA kap. 12", level=4)

    p = doc.add_paragraph()
    add_checkbox(p, "Ingen forberedende unders\u00f8kelser eller dialog med leverand\u00f8rer f\u00f8r konkurransen")
    add_manual(p, " [Bekreft]")

    docx_info_table(doc, [
        ("Forberedende unders\u00f8kelser (jf. FOA \u00a7 12-1):", None),
        ("Leverand\u00f8r(er) som deltok i dialog f\u00f8r konkurransen (jf. FOA \u00a7 12-2):", None),
        ("Avhjelpende tiltak for \u00e5 sikre konkurranse:", None),
    ])

    # § 7-5 — Habilitet
    doc.add_heading("Habilitet, jf. FOA \u00a7 7-5", level=4)

    p = doc.add_paragraph()
    add_checkbox(p, "Ingen habilitetskonflikter identifisert")
    add_manual(p, " [Bekreft]")

    docx_info_table(doc, [
        ("Inhabilitet eller konkurransevridning:", None),
        ("Eventuelle avhjelpende tiltak:", None),
    ])


def _other(doc, procurement):
    doc.add_heading("Andre opplysninger og avslutning", level=3)

    is_cancelled = procurement.get("isCancelled")
    if is_cancelled:
        reason = procurement.get("cancelingReason") or ""
        cancel_str = reason if reason else None
    else:
        cancel_str = "Ikke relevant (konkurransen ble ikke avlyst)."

    docx_info_table(doc, [
        ("Underleverandører (hvilke deler av kontrakten, navn):", None),
        ("Begrunnelse for avlysning, jf. FOA \u00a7 25-4:", cancel_str),
        ("Andre opplysninger, vesentlige forhold eller viktige beslutninger:", None),
    ])


def _data_quality(doc, procurement, activities):
    doc.add_heading("Datakvalitet \u2014 API vs. manuelt", level=3)
    add_instruction(doc, "Intern oversikt \u2014 ikke del av den formelle protokollen. Viser hvilke seksjoner som er fylt ut automatisk fra API-data og hvilke som krever manuell utfylling.")

    submissions = get_activities_by_action(activities, "SUBMIT_BID")
    rejections = get_activities_by_action(activities, "REJECT_PARTICIPATION")
    doffin = get_activities_by_action(activities, "DOFFIN_NOTICE_STATUS_PUBLISHED")
    publish = get_activities_by_action(activities, "PUBLISH_TO_DOFFIN")

    submission_deadline = parse_submission_deadline(procurement)
    conversations = get_activities_by_action(activities, "CONVERSATION_MARKED_COMPLETED")
    post_deadline = []
    if submission_deadline:
        for c in conversations:
            try:
                dt = datetime.fromisoformat((c.get("date") or "").replace("Z", "+00:00"))
                if dt > submission_deadline:
                    post_deadline.append(c)
            except (ValueError, TypeError):
                pass

    rows = [
        # Rammeverk
        ("Generell informasjon", "API", "Komplett"),
        ("Tidspunkt for mottak", "API (SUBMIT_BID)", "Komplett" if submissions else "Ingen tilbud registrert"),
        ("Prosedyretype", "API (procedure)", "Komplett"),
        ("Kunngjøring", f"API {'(DOFFIN_NOTICE)' if doffin or publish else ''}", "Komplett" if doffin or publish else "Trenger bekreftelse"),
        ("Delkontrakter (begrunnelse)", "Manuelt", "Ikke i API"),
        ("Elektronisk kommunikasjon", "Manuelt", "Betinget \u2014 kun ved unntak"),
        # Kvalifisering
        ("Kvalifikasjonsvurdering", "Manuelt", "Ikke i API"),
        ("Utvelgelse", "Manuelt", "Betinget \u2014 kun begrenset/forhandlet"),
        # Avvisning
        ("Avvisning formalfeil \u00a7 24-1", f"API ({len(rejections)} hendelser)" if rejections else "API (ingen hendelser)", "Trenger manuell klassifisering" if rejections else "Trenger bekreftelse"),
        ("Avvisning leverandører \u00a7 24-2", f"API ({len(rejections)} hendelser)" if rejections else "API (ingen hendelser)", "Begrunnelse mangler" if rejections else "Trenger bekreftelse"),
        ("Avvisning tilbud \u00a7\u00a7 24-8/24-9", "API (ingen hendelser)", "Trenger bekreftelse"),
        ("Unormalt lave tilbud \u00a7 24-9", "Manuelt", "Ikke i API"),
        # Ettersending, forhandlinger og dialog
        ("Ettersending/avklaring", f"API ({len(post_deadline)} meldinger etter frist)" if post_deadline else "API (ingen hendelser)", "Innhold mangler" if post_deadline else "Trenger bekreftelse"),
        ("Forhandlinger", "Manuelt", "Betinget \u2014 kun forhandlet prosedyre"),
        ("Dialog", "Manuelt", "Betinget \u2014 kun konkurransepreget dialog"),
        # Tildeling
        ("Tilbud i vurdering", "API (SUBMIT_BID)", "Komplett" if submissions else "Ingen"),
        ("Valgte tilbud + begrunnelse", "Manuelt", "API har kun tenderIds"),
        ("Meddelelsesbrev/karens", "Manuelt", "Kun flag, ikke datoer"),
        ("Rammeavtale", "API", "Betinget"),
        # Avslutning
        ("Markedsdialog/inhabilitet", "Manuelt", "Ikke i API"),
        ("Underleverandører", "Manuelt", "Ikke i API"),
        ("Avlysning", "API (isCancelled)", "Betinget"),
    ]
    docx_add_table(doc, ["Seksjon", "Kilde", "Merknad"], rows)


# -- Main generator ----------------------------------------------------------

def generate_protokoll_docx(procurement: dict, activities: list[dict]) -> DocxDocument:
    """Generate a complete anskaffelsesprotokoll as a Word document (Del III)."""
    procedure = procurement.get("procedure") or ""
    seq_id = procurement.get("sequenceId") or procurement.get("name") or "Ukjent"
    today = datetime.now().strftime("%d.%m.%Y")

    doc = DocxDocument()
    docx_setup(doc)
    org_lookup = build_org_lookup(activities)

    doc.add_heading(f"ANSKAFFELSESPROTOKOLL — {seq_id}", level=1)
    docx_subtitle(doc, seq_id, today)

    doc.add_heading("Rammeverk", level=2)
    _general_info(doc, procurement, activities, org_lookup)
    _procedure(doc, procurement, activities)

    doc.add_heading("Kvalifisering", level=2)
    _preliminary_qualification(doc, procedure)
    _qualification(doc)
    _supplier_selection(doc, procedure, activities)

    doc.add_heading("Avvisning", level=2)
    _formal_rejection(doc, activities, org_lookup)
    _supplier_rejection(doc, activities, org_lookup)
    _bid_rejection(doc)

    doc.add_heading("Ettersending, forhandlinger og dialog", level=2)
    _clarification(doc, procurement, activities, org_lookup)
    _negotiations(doc, procedure)
    _dialog(doc, procedure)

    doc.add_heading("Tildeling", level=2)
    _bids_in_evaluation(doc, activities, org_lookup)
    _award(doc, procurement, activities)
    _framework_agreement(doc, procurement)

    doc.add_heading("Avslutning", level=2)
    _market_dialogue_and_conflicts(doc)
    _other(doc, procurement)
    _data_quality(doc, procurement, activities)

    return doc
