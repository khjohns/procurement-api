"""Word document generator for Del II (nasjonal, under EØS) anskaffelsesprotokoll."""

from __future__ import annotations

import re
from datetime import datetime

from docx import Document as DocxDocument

from .common import (
    ACTION_AWARDING_PARTICIPANTS,
    ACTION_DOFFIN_NOTICE_STATUS_PUBLISHED,
    ACTION_PUBLISH_TO_DOFFIN,
    ACTION_QUALIFYING_PARTICIPANTS,
    ACTION_REJECT_PARTICIPATION,
    ACTION_SUBMIT_BID,
    DEL2_PROCEDURE_MAP,
    TIMELINE_AWARD_DECISION,
    TIMELINE_SUBMISSION,
    build_org_lookup,
    filter_post_deadline_conversations,
    fmt_currency,
    fmt_date,
    fmt_datetime,
    get_activities_by_action,
    get_org_name,
    get_timeline_date,
    parse_announcement,
    strip_html,
)
from .docx_helpers import (
    add_checkbox,
    add_instruction,
    add_manual,
    add_text_box,
    docx_add_table,
    docx_add_table_with_manual,
    docx_info_table,
    docx_setup,
    docx_subtitle,
)

from eforms_labels import get_label

# Reuse shared sections from Del III
from .docx_del3 import (
    _award_criteria,
    _bids_in_evaluation,
    _cancellation,
    _contract_modifications,
    _framework_agreement,
)


def _general_info(doc, procurement, activities, org_lookup):
    doc.add_heading("Generell informasjon", level=3)

    seq_id = procurement.get("sequenceId") or ""
    ext_id = procurement.get("externalId") or ""
    sak_ref = seq_id
    if ext_id:
        sak_ref += f" (ekstern ref: {ext_id})"

    procurer = procurement.get("about_procurer") or {}
    procurer_name = procurer.get("name") or ""
    national_id = procurer.get("national_id") or ""
    procurer_str = procurer_name
    if national_id:
        procurer_str += f" (org.nr. {national_id})"

    contact = procurer.get("contact_person") or ""

    name = strip_html(procurement.get("name") or "")
    description = strip_html(procurement.get("description") or "")
    description = re.sub(r"\s*\n\s*", " ", description)
    desc_str = name
    if description and description.lower() != name.lower():
        desc_str += f". {description}"

    estimated_value = procurement.get("estimated_value")
    currency = procurement.get("currency") or "NOK"
    value_str = fmt_currency(estimated_value, currency)

    duration_months = procurement.get("duration_months")
    duration = procurement.get("duration")
    if duration_months:
        duration_str = f"{duration_months} måneder"
    elif duration:
        duration_str = str(duration)
    else:
        duration_str = None

    submission_date = get_timeline_date(procurement, TIMELINE_SUBMISSION)
    submission_str = fmt_datetime(submission_date) if submission_date else None

    docx_info_table(
        doc,
        [
            ("Konkurransens saksnummer:", sak_ref),
            ("Oppdragsgiver(e):", procurer_str),
            ("Protokollfører/saksbehandler:", contact or None),
            ("Beskrivelse av anskaffelsen:", desc_str),
            ("Kontraktens anslåtte verdi på kunngjøringstidspunktet:", value_str),
            ("Kontraktens varighet:", duration_str),
            ("Frist for innlevering av tilbud:", submission_str),
        ],
    )

    submissions = get_activities_by_action(activities, ACTION_SUBMIT_BID)
    p = doc.add_paragraph()
    p.add_run("Tidspunkt for mottak av tilbud:").bold = True

    if submissions:
        rows = []
        for s in submissions:
            org_name = get_org_name(s, org_lookup)
            date = fmt_datetime(s.get("date"))
            rows.append((org_name, date))
        docx_add_table(doc, ["Leverandørens navn", "Tidspunkt for mottak"], rows)
    else:
        docx_add_table(
            doc,
            ["Leverandørens navn", "Tidspunkt for mottak"],
            [("Ingen tilbud registrert", "")],
        )


def _procedure(doc, procurement, activities, eforms=None):
    doc.add_heading("Anskaffelsesprosedyre", level=3)

    procedure = procurement.get("procedure") or ""
    selected = DEL2_PROCEDURE_MAP.get(procedure, procedure)
    p = doc.add_paragraph()
    p.add_run("Prosedyre: ").bold = True
    p.add_run(selected)

    if eforms:
        nature = eforms.get("contract_nature")
        if nature:
            p2 = doc.add_paragraph()
            p2.add_run("Kontraktstype: ").bold = True
            p2.add_run(get_label("contract-nature", nature, nature))

    # Kunngjøring
    announcement_date, doffin_ref, ted_ref = parse_announcement(activities)

    if announcement_date:
        parts = [f"Kunngjort {announcement_date} i DOFFIN"]
        if doffin_ref:
            parts.append(f"ref. {doffin_ref}")
        if ted_ref:
            parts.append(f"TED {ted_ref}")
        kunngj_str = ", ".join(parts)
    else:
        pub_date = procurement.get("publicationDate")
        kunngj_str = f"Kunngjort {fmt_date(pub_date)}" if pub_date else None

    docx_info_table(
        doc,
        [
            ("Kunngjøring:", kunngj_str),
        ],
    )

    # Betingede lovkrav (§ 10-5 nr. 3 og nr. 14)
    p = doc.add_paragraph()
    add_checkbox(p, "Unntak fra krav om elektronisk kommunikasjon ved mottak av tilbud")
    add_manual(p, " [Fyll inn begrunnelse, jf. FOA \u00a7 10-5 (2) nr. 3]")

    p = doc.add_paragraph()
    add_checkbox(
        p,
        "Konkurransen er reservert for ideelle organisasjoner (helse-/sosialtjenester)",
    )
    add_manual(p, " [Fyll inn begrunnelse, jf. FOA \u00a7 10-5 (2) nr. 14]")


def _dialog(doc, procurement, activities):
    """Dialog og/eller forhandlinger, jf. FOA § 9-3."""
    doc.add_heading("Dialog og/eller forhandlinger, jf. FOA \u00a7 9-3", level=3)

    p = doc.add_paragraph()
    p.add_run(
        "Hvis relevant, begrunnelse for å fravike opplysningene i anskaffelsesdokumentene om planlagt dialog:"
    ).bold = True
    p2 = doc.add_paragraph()
    add_manual(p2)

    p = doc.add_paragraph()
    add_checkbox(p, "Ingen dialog eller forhandlinger gjennomført")
    add_manual(p, " [Bekreft]")

    docx_add_table_with_manual(
        doc,
        [
            "Leverandørens navn",
            "Dato",
            "Begrunnelse for hvorfor leverandøren er valgt ut til dialog",
        ],
        [(None, None, None)],
    )

    docx_info_table(
        doc,
        [
            (
                "Inhabilitet/konkurransevridning som følge av dialog, og avhjelpende tiltak:",
                None,
            ),
        ],
    )


def _formal_rejection(doc, activities, org_lookup):
    """Avvisning formalfeil, jf. FOA § 9-4."""
    doc.add_heading("Avvisning på grunn av formalfeil, jf. FOA \u00a7 9-4", level=3)

    rejections = get_activities_by_action(activities, ACTION_REJECT_PARTICIPATION)
    if not rejections:
        p = doc.add_paragraph()
        add_checkbox(p, "Ingen avvist på grunn av formalfeil")
        add_manual(p, " [Bekreft]")
        return

    p = doc.add_paragraph()
    add_manual(
        p,
        "[Avgjør hvilke avvisninger som gjelder formalfeil (\u00a7 9-4) vs. kvalifikasjon (\u00a7 9-5)]",
    )

    p2 = doc.add_paragraph(
        "Merk: Avvisningene nedenfor vises i begge avvisningsseksjonene. "
        "Slett de som ikke gjelder denne paragrafen."
    )
    p2.runs[0].italic = True

    rows = []
    for r in rejections:
        org_name = get_org_name(r, org_lookup)
        date = fmt_date(r.get("date"))
        rows.append((org_name, None, date))
    docx_add_table_with_manual(
        doc,
        ["Leverandørens navn", "Begrunnelsen for avvisningen", "Begrunnelse ble sendt"],
        rows,
    )


def _qualification(doc, eforms=None):
    """Three-tier qualification (full, preliminary § 8-10, chosen supplier)."""
    doc.add_heading("Kvalifikasjonsvurdering", level=3)

    if eforms:
        sel = eforms.get("selection_criteria") or []
        if sel:
            doc.add_paragraph("Kvalifikasjonskrav fra kunngjøringen:")
            rows = [
                (
                    get_label("selection-criterion", s.get("type_code") or "", s.get("type_code") or ""),
                    s.get("description") or "",
                )
                for s in sel
            ]
            docx_add_table(doc, ["Type", "Beskrivelse"], rows)

    # Tier 1: Full qualification
    doc.add_heading("Full kvalifikasjonsvurdering", level=4)
    add_instruction(
        doc,
        "Strykes hvis ikke relevant. Benyttes hvis det ikke kreves egenerklæring som foreløpig dokumentasjonsbevis.",
    )
    p = doc.add_paragraph()
    add_manual(
        p,
        "[Kvalifikasjonskrav og -vurdering. Fyll inn basert på konkurransegrunnlaget.]",
    )

    docx_info_table(
        doc,
        [
            ("Leverandører som er kvalifisert:", None),
            ("Begrunnelse for leverandører med skatte-/avgiftsrestanser:", None),
        ],
    )

    # Tier 2: Preliminary qualification via self-declaration
    doc.add_heading("Foreløpig kvalifikasjonsvurdering", level=4)
    add_instruction(
        doc,
        "Strykes hvis ikke relevant. Benyttes hvis det kreves egenerklæring som foreløpig dokumentasjonsbevis, jf. FOA \u00a7 8-10 (1).",
    )

    docx_info_table(
        doc,
        [
            ("Leverandører som er kvalifisert:", None),
        ],
    )

    # Tier 3: Qualification of chosen supplier(s)
    doc.add_heading("Kvalifikasjonsvurdering av valgte leverand\u00f8r(er)", level=4)
    add_instruction(
        doc,
        "Strykes hvis ikke relevant. Benyttes når valgt leverandør skal levere oppdaterte dokumentasjonsbevis, jf. FOA \u00a7 8-10 (2).",
    )

    docx_info_table(
        doc,
        [
            ("Leverandør(er) som er kvalifisert:", None),
            ("Begrunnelse for leverandør med skatte-/avgiftsrestanser:", None),
        ],
    )


def _supplier_rejection(doc, activities, org_lookup):
    """Leverandører avvist, jf. FOA § 9-5."""
    doc.add_heading("Leverandører som er avvist, jf. FOA \u00a7 9-5", level=3)

    rejections = get_activities_by_action(activities, ACTION_REJECT_PARTICIPATION)
    if not rejections:
        p = doc.add_paragraph()
        add_checkbox(p, "Ingen leverandører avvist")
        add_manual(p, " [Bekreft]")
        return

    doc.add_paragraph("Følgende leverandører ble avvist:")

    p2 = doc.add_paragraph(
        "Merk: Avvisningene nedenfor vises i begge avvisningsseksjonene. "
        "Slett de som ikke gjelder denne paragrafen."
    )
    p2.runs[0].italic = True

    rows = []
    for r in rejections:
        org_name = get_org_name(r, org_lookup)
        date = fmt_date(r.get("date"))
        rows.append((org_name, None, date))
    docx_add_table_with_manual(
        doc,
        ["Leverandørens navn", "Begrunnelsen for avvisningen", "Begrunnelse ble sendt"],
        rows,
    )


def _supplier_selection(doc, procedure, activities):
    """Utvelgelse — only for begrenset tilbudskonkurranse."""
    doc.add_heading("Utvelgelse av leverandører", level=3)

    if procedure != "Limited":
        add_instruction(
            doc,
            "Strykes hvis ikke relevant. Gjelder kun ved begrenset tilbudskonkurranse.",
        )
        doc.add_paragraph("Ikke relevant (åpen tilbudskonkurranse).")
        return

    add_instruction(doc, "Gjelder kun ved begrenset tilbudskonkurranse.")
    qualifying = get_activities_by_action(activities, ACTION_QUALIFYING_PARTICIPANTS)
    p = doc.add_paragraph()
    add_manual(p, "[Fyll inn begrunnelse for utvelgelse per leverandør]")
    if qualifying:
        rows = []
        for q in qualifying:
            desc = q.get("description") or {}
            tenders_ids = desc.get("tendersIds") or []
            for tid in tenders_ids:
                rows.append((f"Leverandør (tender {tid})", None))
        if rows:
            docx_add_table_with_manual(
                doc, ["Leverandørens navn", "Begrunnelse for utvelgelse"], rows
            )
    else:
        docx_add_table_with_manual(
            doc, ["Leverandørens navn", "Begrunnelse for utvelgelse"], [(None, None)]
        )


def _bid_rejection(doc):
    """Tilbud avvist, jf. FOA § 9-6."""
    doc.add_heading("Tilbud som er avvist, jf. FOA \u00a7 9-6", level=3)
    p = doc.add_paragraph()
    add_checkbox(p, "Ingen tilbud avvist")
    add_manual(p, " [Bekreft]")


def _clarification(doc, procurement, activities, org_lookup):
    """Ettersending og avklaring (basert på grunnleggende prinsipper)."""
    doc.add_heading("Ettersending og avklaring", level=3)

    post_deadline_convs, conversations = filter_post_deadline_conversations(
        procurement, activities
    )

    if not post_deadline_convs:
        p = doc.add_paragraph()
        add_checkbox(p, "Det ble ikke foretatt avklaringer eller ettersendinger")
        if not conversations:
            add_manual(p, " [Bekreft \u2014 ingen avklaringshendelser i API]")
        else:
            add_manual(
                p,
                " [Bekreft \u2014 meldinger finnes, men alle er f\u00f8r tilbudsfrist (Q&A)]",
            )
        return

    doc.add_paragraph(
        "F\u00f8lgende avklaringer/ettersendinger ble gjennomf\u00f8rt etter tilbudsfrist:"
    )
    rows = []
    for c in post_deadline_convs:
        org_name = get_org_name(c, org_lookup)
        date = fmt_date(c.get("date"))
        desc = c.get("description") or {}
        title = desc.get("conversationTitle") or ""
        how = f"Melding i KGV: \u00ab{title}\u00bb" if title else "Melding i KGV"
        rows.append((org_name, date, how))
    docx_add_table(doc, ["Leverandørens navn", "Dato", "Hvordan?"], rows)


def _award(doc, procurement, activities):
    """Valgt tilbud med begrunnelse og kontraktsverdi."""
    doc.add_heading("Det (de) valgte tilbud med begrunnelse og kontraktsverdi", level=3)

    add_text_box(
        doc,
        "[Fyll inn navn på valgt leverandør, tildelingsbegrunnelse og kontraktsverdi. Begrunnelsen skal inneholde tilstrekkelig informasjon til at øvrige leverandører kan vurdere om valget er saklig og forsvarlig.]",
    )

    total_value = procurement.get("contracts_total_value_amount")
    estimated = procurement.get("estimated_value")
    currency = procurement.get("currency") or "NOK"
    if total_value:
        value_str = fmt_currency(total_value, currency)
    elif estimated:
        value_str = f"{fmt_currency(estimated, currency)} (estimert verdi)"
    else:
        value_str = None

    award_date = get_timeline_date(procurement, TIMELINE_AWARD_DECISION)
    if not award_date:
        award_activities = get_activities_by_action(
            activities, ACTION_AWARDING_PARTICIPANTS
        )
        if award_activities:
            award_date = award_activities[0].get("date")
    award_str = fmt_date(award_date) if award_date else None

    docx_info_table(
        doc,
        [
            ("Kontraktsverdi:", value_str),
            ("Tildelingsbeslutning:", award_str),
        ],
    )


def _award_notification(doc, procurement):
    """Meddelelse om tildeling og klagefrist (not karensperiode)."""
    doc.add_heading(
        "Meddelelse om tildeling og klagefrist f\u00f8r inng\u00e5else av kontrakt",
        level=3,
    )

    award_letters = procurement.get("areAwardLettersSent")
    docx_info_table(
        doc,
        [
            (
                "Meddelelsesbrev sendt:",
                "Sendt (dato ikke tilgjengelig i API)" if award_letters else None,
            ),
            ("Klagefrist:", None),
            ("Eventuelle klager:", None),
            ("Resultat av klage:", None),
        ],
    )


def _market_dialogue_and_conflicts(doc):
    """Dialog med markedet og habilitet, jf. FOA kap. 8 og § 7-5."""
    doc.add_heading("Dialog med markedet og habilitet", level=3)

    # Kap. 8 — Planlegging og kunngjøring (§§ 8-1 og 8-2)
    doc.add_heading(
        "Dialog med markedet f\u00f8r konkurranse, jf. FOA \u00a7\u00a7 8-1 og 8-2",
        level=4,
    )

    p = doc.add_paragraph()
    add_checkbox(
        p,
        "Ingen forberedende unders\u00f8kelser eller dialog med leverand\u00f8rer f\u00f8r konkurransen",
    )
    add_manual(p, " [Bekreft]")

    docx_info_table(
        doc,
        [
            ("Forberedende unders\u00f8kelser (jf. FOA \u00a7 8-1):", None),
            (
                "Leverand\u00f8r(er) som deltok i dialog f\u00f8r konkurransen (jf. FOA \u00a7 8-2):",
                None,
            ),
            ("Avhjelpende tiltak for \u00e5 sikre konkurranse:", None),
        ],
    )

    # § 7-5 — Habilitet
    doc.add_heading("Habilitet, jf. FOA \u00a7 7-5", level=4)

    p = doc.add_paragraph()
    add_checkbox(p, "Ingen habilitetskonflikter identifisert")
    add_manual(p, " [Bekreft]")

    docx_info_table(
        doc,
        [
            ("Inhabilitet eller konkurransevridning:", None),
            ("Eventuelle avhjelpende tiltak:", None),
        ],
    )


def _other(doc, procurement):
    """Andre opplysninger — § 10-4 for avlysning."""
    doc.add_heading("Andre opplysninger og avslutning", level=3)

    is_cancelled = procurement.get("isCancelled")
    if is_cancelled:
        reason = procurement.get("cancelingReason") or ""
        cancel_str = reason if reason else None
    else:
        cancel_str = "Ikke relevant (konkurransen ble ikke avlyst)."

    docx_info_table(
        doc,
        [
            ("Underleverandører (hvilke deler av kontrakten, navn):", None),
            ("Begrunnelse for avlysning, jf. FOA \u00a7 10-4 (1):", cancel_str),
            (
                "Andre opplysninger, vesentlige forhold eller viktige beslutninger:",
                None,
            ),
        ],
    )


def _data_quality(doc, procurement, activities, eforms=None):
    doc.add_heading("Datakvalitet \u2014 API vs. manuelt", level=3)
    add_instruction(
        doc,
        "Intern oversikt \u2014 ikke del av den formelle protokollen. Viser hvilke seksjoner som er fylt ut automatisk fra API-data og hvilke som krever manuell utfylling.",
    )

    submissions = get_activities_by_action(activities, ACTION_SUBMIT_BID)
    rejections = get_activities_by_action(activities, ACTION_REJECT_PARTICIPATION)
    doffin = get_activities_by_action(activities, ACTION_DOFFIN_NOTICE_STATUS_PUBLISHED)
    publish = get_activities_by_action(activities, ACTION_PUBLISH_TO_DOFFIN)

    post_deadline, _ = filter_post_deadline_conversations(procurement, activities)

    has_eforms = eforms is not None
    eforms_ac = len((eforms or {}).get("award_criteria") or [])
    eforms_sc = len((eforms or {}).get("selection_criteria") or [])

    rows = [
        # Rammeverk
        ("Generell informasjon", "API", "Komplett"),
        (
            "Tidspunkt for mottak",
            "API (SUBMIT_BID)",
            "Komplett" if submissions else "Ingen tilbud registrert",
        ),
        ("Prosedyretype", "API (procedure)", "Komplett"),
        (
            "Kunngjøring",
            f"API {'(DOFFIN_NOTICE)' if doffin or publish else ''}",
            "Komplett" if doffin or publish else "Trenger bekreftelse",
        ),
        ("Elektronisk kommunikasjon", "Manuelt", "Betinget \u2014 kun ved unntak"),
        (
            "Reservasjon ideelle org.",
            "Manuelt",
            "Betinget \u2014 kun helse-/sosialtjenester",
        ),
        # Dialog og avklaring
        ("Dialog/forhandlinger \u00a7 9-3", "Manuelt", "Ikke i API"),
        (
            "Ettersending/avklaring",
            f"API ({len(post_deadline)} meldinger etter frist)"
            if post_deadline
            else "API (ingen hendelser)",
            "Innhold mangler" if post_deadline else "Trenger bekreftelse",
        ),
        # Kvalifisering
        (
            "Tildelingskriterier",
            f"eForms ({eforms_ac} kriterier)"
            if has_eforms and eforms_ac
            else "Manuelt",
            "Komplett" if eforms_ac else "Ikke i API",
        ),
        (
            "Kvalifikasjonskrav",
            f"eForms ({eforms_sc} krav)" if has_eforms and eforms_sc else "Manuelt",
            "Komplett"
            if eforms_sc
            else "Ikke i eForms (ligger i konkurransegrunnlaget)",
        ),
        # Avvisning
        (
            "Avvisning formalfeil \u00a7 9-4",
            f"API ({len(rejections)} hendelser)"
            if rejections
            else "API (ingen hendelser)",
            "Trenger manuell klassifisering" if rejections else "Trenger bekreftelse",
        ),
        (
            "Avvisning leverandører \u00a7 9-5",
            f"API ({len(rejections)} hendelser)"
            if rejections
            else "API (ingen hendelser)",
            "Begrunnelse mangler" if rejections else "Trenger bekreftelse",
        ),
        ("Avvisning tilbud \u00a7 9-6", "API (ingen hendelser)", "Trenger bekreftelse"),
        # Tildeling
        (
            "Tilbud i vurdering",
            "API (SUBMIT_BID)",
            "Komplett" if submissions else "Ingen",
        ),
        ("Valgte tilbud + begrunnelse", "Manuelt", "API har kun tenderIds"),
        ("Meddelelsesbrev/klagefrist", "Manuelt", "Kun flag, ikke datoer"),
        ("Rammeavtale", "API (framework)", "Betinget \u2014 kun rammeavtaler"),
        # Avslutning
        ("Dialog med markedet \u00a7\u00a7 8-1, 8-2", "Manuelt", "Ikke i API"),
        ("Habilitet \u00a7 7-5", "Manuelt", "Ikke i API"),
        ("Underleverandører", "Manuelt", "Ikke i API"),
    ]
    docx_add_table(doc, ["Seksjon", "Kilde", "Merknad"], rows)


# -- Main generator ----------------------------------------------------------


def generate_protokoll_docx_del2(
    procurement: dict, activities: list[dict], eforms=None
) -> DocxDocument:
    """Generate anskaffelsesprotokoll for Del II (nasjonal, under EØS) as Word document."""
    procedure = procurement.get("procedure") or ""
    seq_id = procurement.get("sequenceId") or procurement.get("name") or "Ukjent"
    today = datetime.now().strftime("%d.%m.%Y")

    doc = DocxDocument()
    docx_setup(doc)
    org_lookup = build_org_lookup(activities)

    doc.add_heading(
        f"ANSKAFFELSESPROTOKOLL for anskaffelser etter forskriften del II — {seq_id}",
        level=1,
    )
    docx_subtitle(doc, seq_id, today)

    doc.add_heading("Rammeverk", level=2)
    _general_info(doc, procurement, activities, org_lookup)
    _procedure(doc, procurement, activities, eforms)

    doc.add_heading("Dialog og avklaring", level=2)
    _dialog(doc, procurement, activities)
    _clarification(doc, procurement, activities, org_lookup)

    doc.add_heading("Kvalifisering", level=2)
    _qualification(doc, eforms)
    _supplier_selection(doc, procedure, activities)

    doc.add_heading("Avvisning", level=2)
    _formal_rejection(doc, activities, org_lookup)
    _supplier_rejection(doc, activities, org_lookup)
    _bid_rejection(doc)

    is_cancelled = procurement.get("isCancelled", False)

    if is_cancelled:
        doc.add_heading("Avlysning", level=2)
        _cancellation(doc, procurement)
    else:
        doc.add_heading("Tildeling", level=2)
        _award_criteria(doc, eforms)
        _bids_in_evaluation(doc, activities, org_lookup)
        _award(doc, procurement, activities)
        _award_notification(doc, procurement)
        _framework_agreement(doc, procurement, eforms)

        doc.add_heading("Kontraktsendringer", level=2)
        _contract_modifications(doc)

    doc.add_heading("Avslutning", level=2)
    _market_dialogue_and_conflicts(doc)
    _other(doc, procurement)
    _data_quality(doc, procurement, activities, eforms)

    return doc
