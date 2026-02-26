"""Shared Word document helpers for protokoll generators."""

from __future__ import annotations

try:
    from docx.shared import Pt, RGBColor, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False


def docx_setup(doc):
    """Configure document margins, styles, and page footer."""
    # Narrower margins (2.0 cm)
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Heading 2: add space before for section separation
    h2_style = doc.styles["Heading 2"]
    h2_style.paragraph_format.space_before = Pt(18)
    h2_style.paragraph_format.space_after = Pt(6)

    # Page numbers in footer
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = 1  # CENTER
        run = p.add_run()
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), r" PAGE ")
        run._element.append(fld)
        run2 = p.add_run(" / ")
        run2.font.size = Pt(8)
        run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run3 = p.add_run()
        fld2 = OxmlElement("w:fldSimple")
        fld2.set(qn("w:instr"), r" NUMPAGES ")
        run3._element.append(fld2)
        for r in [run, run3]:
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def docx_shade_header_row(table):
    """Apply light grey background to the first row of a table."""
    for cell in table.rows[0].cells:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), "F2F2F2")
        cell._element.get_or_add_tcPr().append(shading)


def docx_set_col_widths(table, widths_cm):
    """Set column widths in cm."""
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            if i < len(row.cells):
                row.cells[i].width = Cm(w)


def add_manual(paragraph, text="[Fyll inn]"):
    """Add a yellow-highlighted, italic manual marker to a paragraph."""
    run = paragraph.add_run(text)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x66, 0x00)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "FFFF00")
    run._element.get_or_add_rPr().append(shading)
    return run


def docx_add_table(doc, headers, rows):
    """Add a table with bold headers, grey header row, and bordered cells."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_data in enumerate(row_data):
            table.rows[r_idx + 1].cells[c_idx].text = str(cell_data)
    docx_shade_header_row(table)
    return table


def docx_add_table_with_manual(doc, headers, rows):
    """Add a table where cells containing None get a yellow manual marker."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_data in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if cell_data is None:
                cell.text = ""
                add_manual(cell.paragraphs[0])
            else:
                cell.text = str(cell_data)
    docx_shade_header_row(table)
    return table


def docx_info_table(doc, info_rows):
    """Label-value table with 6.0/11.0 cm column widths."""
    table = doc.add_table(rows=len(info_rows), cols=2)
    table.style = "Table Grid"
    docx_set_col_widths(table, [6.0, 11.0])
    for i, (label, val) in enumerate(info_rows):
        cell_label = table.rows[i].cells[0]
        cell_label.text = ""
        run = cell_label.paragraphs[0].add_run(label)
        run.bold = True
        cell_val = table.rows[i].cells[1]
        if val is None:
            cell_val.text = ""
            add_manual(cell_val.paragraphs[0])
        else:
            cell_val.text = str(val)
    return table


def docx_subtitle(doc, seq_id, today):
    """Add the subtitle paragraph with generation date and manual marker legend."""
    p = doc.add_paragraph()
    run = p.add_run(f"Generert fra API-data {today}. ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run2 = p.add_run("Felter med gul bakgrunn krever manuell utfylling.")
    run2.font.size = Pt(9)
    run2.font.italic = True
    run2.font.color.rgb = RGBColor(0x99, 0x66, 0x00)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "FFFF00")
    run2._element.get_or_add_rPr().append(shading)
