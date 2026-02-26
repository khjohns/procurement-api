"""Generate anskaffelsesprotokoll from Artifik API data.

Standalone CLI — fetches secrets from GCP Secret Manager, talks directly to
Artifik API via ArtifikClient. No MCP dependency.

Usage:
    python3 src/protokoll_generator.py              # interactive selection
    python3 src/protokoll_generator.py --id 1795    # specific procurement
    python3 src/protokoll_generator.py --list       # just list procurements
"""

from __future__ import annotations

import html
import itertools
import os
import re
import subprocess
import sys
import threading
import time
from datetime import datetime
from pathlib import Path
from typing import Any

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False

# Add src/ to path so we can import ArtifikClient
sys.path.insert(0, str(Path(__file__).resolve().parent))

from app.client import ArtifikClient  # noqa: E402

GCP_PROJECT = "procurement-mcp"

# -- Terminal formatting -----------------------------------------------------

_COLOR = hasattr(sys.stderr, "isatty") and sys.stderr.isatty()


def _style(text: str, code: str) -> str:
    return f"\033[{code}m{text}\033[0m" if _COLOR else text


def _bold(text: str) -> str:
    return _style(text, "1")


def _dim(text: str) -> str:
    return _style(text, "2")


def _green(text: str) -> str:
    return _style(text, "32")


def _yellow(text: str) -> str:
    return _style(text, "33")


def _red(text: str) -> str:
    return _style(text, "31")


def _cyan(text: str) -> str:
    return _style(text, "36")


class _Spinner:
    """Minimal terminal spinner for long-running operations."""

    _FRAMES = ["   ", ".  ", ".. ", "..."]

    def __init__(self, message: str) -> None:
        self._message = message
        self._stop = threading.Event()
        self._thread: threading.Thread | None = None

    def __enter__(self) -> _Spinner:
        self._stop.clear()
        self._thread = threading.Thread(target=self._spin, daemon=True)
        self._thread.start()
        return self

    def __exit__(self, *_: Any) -> None:
        self._stop.set()
        if self._thread:
            self._thread.join()
        # Clear the spinner line
        print("\r\033[K", end="", file=sys.stderr)

    def _spin(self) -> None:
        for frame in itertools.cycle(self._FRAMES):
            if self._stop.is_set():
                break
            print(
                f"\r  {_dim(self._message)}{_dim(frame)}",
                end="",
                file=sys.stderr,
                flush=True,
            )
            time.sleep(0.3)


def _step(number: int, total: int, text: str) -> None:
    label = _dim(f"[{number}/{total}]")
    print(f"  {label} {text}", file=sys.stderr)


def _ok(text: str) -> None:
    print(f"  {_green('OK')} {text}", file=sys.stderr)


def _warn(text: str) -> None:
    print(f"  {_yellow('!')} {text}", file=sys.stderr)


# -- Procedure mapping -------------------------------------------------------

PROCEDURE_MAP = {
    "Open": "Åpen anbudskonkurranse",
    "Limited": "Begrenset anbudskonkurranse",
    "Competitive negotiated": "Konkurranse med forhandling etter forutgående kunngjøring",
    "Competitive dialogue": "Konkurransepreget dialog",
    "Innovation partnership": "Konkurranse om innovasjonspartnerskap",
    "Negotiated without publication": "Konkurranse med forhandling uten forutgående kunngjøring",
    "Direct award": "Anskaffelse uten konkurranse",
}

ALL_PROCEDURES = [
    "Åpen anbudskonkurranse",
    "Begrenset anbudskonkurranse",
    "Konkurranse med forhandling etter forutgående kunngjøring",
    "Konkurransepreget dialog",
    "Konkurranse om innovasjonspartnerskap",
    "Konkurranse med forhandling uten forutgående kunngjøring",
    "Anskaffelse uten konkurranse",
]

DEL2_PROCEDURE_MAP = {
    "Open": "Åpen tilbudskonkurranse",
    "Limited": "Begrenset tilbudskonkurranse",
}

ALL_DEL2_PROCEDURES = [
    "Åpen tilbudskonkurranse",
    "Begrenset tilbudskonkurranse",
]


# -- GCP Secret Manager -----------------------------------------------------

def _fetch_secret(name: str) -> str:
    """Fetch a secret from GCP Secret Manager via gcloud CLI."""
    result = subprocess.run(
        [
            "gcloud", "secrets", "versions", "access", "latest",
            f"--secret={name}",
            f"--project={GCP_PROJECT}",
        ],
        capture_output=True,
        text=True,
    )
    if result.returncode != 0:
        stderr = result.stderr.strip()
        _die(
            f"Kunne ikke hente secret {_bold(name)} fra GCP Secret Manager.\n"
            f"  {_dim(stderr)}\n\n"
            f"  Kjør:  {_cyan('gcloud auth login')}\n"
            f"  Sjekk: {_cyan(f'gcloud config set project {GCP_PROJECT}')}"
        )
    return result.stdout.strip()


def _get_client() -> ArtifikClient:
    """Create ArtifikClient with credentials from GCP Secret Manager."""
    with _Spinner("Henter API-nøkler fra GCP Secret Manager"):
        api_id = _fetch_secret("vendor-api-id")
        api_key = _fetch_secret("vendor-api-key")
    _ok("API-nøkler hentet")
    return ArtifikClient(client_id=api_id, client_secret=api_key)


# -- Helpers -----------------------------------------------------------------

def _die(msg: str) -> None:
    print(f"\n  {_red('Feil:')} {msg}\n", file=sys.stderr)
    sys.exit(1)


def _fmt_datetime(iso_str: str | None) -> str:
    """Format ISO datetime to 'DD.MM.YYYY, kl. HH:MM'."""
    if not iso_str:
        return ""
    try:
        dt = datetime.fromisoformat(iso_str.replace("Z", "+00:00"))
        return dt.strftime("%d.%m.%Y, kl. %H:%M")
    except (ValueError, TypeError):
        return str(iso_str)


def _fmt_date(iso_str: str | None) -> str:
    """Format ISO datetime to 'DD.MM.YYYY'."""
    if not iso_str:
        return ""
    try:
        dt = datetime.fromisoformat(iso_str.replace("Z", "+00:00"))
        return dt.strftime("%d.%m.%Y")
    except (ValueError, TypeError):
        return str(iso_str)


def _fmt_currency(value: Any, currency: str = "NOK") -> str:
    """Format a numeric value as currency."""
    if value is None:
        return ""
    try:
        num = float(value)
        if num == int(num):
            return f"{int(num):,} {currency}".replace(",", " ")
        return f"{num:,.2f} {currency}".replace(",", " ")
    except (ValueError, TypeError):
        return str(value)


def _get_timeline_date(procurement: dict, timeline_type: str) -> str | None:
    """Get a date from the procurement timeline by type."""
    timeline = procurement.get("timeline") or {}
    for entry in timeline.values() if isinstance(timeline, dict) else timeline:
        item = entry if isinstance(entry, dict) else {}
        if item.get("type") == timeline_type:
            return item.get("date")
    return None


def _get_activities_by_action(activities: list[dict], action: str) -> list[dict]:
    """Filter activities by action type, sorted by date."""
    result = [a for a in activities if a.get("action") == action]
    result.sort(key=lambda a: a.get("date") or "")
    return result


def _parse_submission_deadline(procurement: dict) -> datetime | None:
    """Parse the submission deadline from timeline."""
    date_str = _get_timeline_date(procurement, "submission")
    if not date_str:
        return None
    try:
        return datetime.fromisoformat(date_str.replace("Z", "+00:00"))
    except (ValueError, TypeError):
        return None


# -- Procurement listing & selection -----------------------------------------

def _is_mature(procurement: dict) -> bool:
    """Check if procurement is past submission deadline and not a template."""
    if procurement.get("isTemplate"):
        return False
    if procurement.get("isCancelled"):
        return False
    deadline = _parse_submission_deadline(procurement)
    if not deadline:
        return False
    return datetime.now(deadline.tzinfo) > deadline


def _richness(p: dict) -> int:
    """Score how much data a procurement object contains."""
    return sum(1 for v in p.values() if v is not None and v != "" and v != [] and v != {})


def _dedup_by_sequence_id(procs: list[dict]) -> list[dict]:
    """Deduplicate procurements by sequenceId, keeping the richest one."""
    best: dict[str, dict] = {}
    for p in procs:
        seq = p.get("sequenceId") or str(p.get("id"))
        existing = best.get(seq)
        if existing is None or _richness(p) > _richness(existing):
            best[seq] = p
    return list(best.values())


def _list_procurements(client: ArtifikClient) -> list[dict]:
    """Fetch and filter mature procurements."""
    with _Spinner("Henter anskaffelser fra Artifik"):
        all_procs = client.list_procurements()
    mature = [p for p in all_procs if _is_mature(p)]
    mature = _dedup_by_sequence_id(mature)
    mature.sort(key=lambda p: _get_timeline_date(p, "submission") or "", reverse=True)
    _ok(f"{len(mature)} anskaffelser med passert tilbudsfrist (av {len(all_procs)} totalt)")
    return mature


THRESHOLD_SHORT = {
    "over_eea_threshold_value": "Over EØS",
    "below_eea_threshold_value": "Under EØS",
    "national_threshold": "Nasjonal",
    "below_national_threshold": "Under terskel",
}

PROCEDURE_SHORT = {
    "Open": "Åpen",
    "Limited": "Begrenset",
    "Competitive negotiated": "Forhandl.",
    "Competitive dialogue": "Dialog",
    "Innovation partnership": "Innovasjon",
    "Negotiated without publication": "Utenkunng.",
    "Direct award": "Direkte",
}


def _color_threshold(raw: str, label: str) -> str:
    """Color-code threshold: higher regulation = warmer color."""
    if raw == "over_eea_threshold_value":
        return _style(label, "1;35")  # bold magenta — del III
    if raw == "below_eea_threshold_value":
        return _cyan(label)            # cyan — del II
    return _dim(label)                 # dim — nasjonal/under terskel


def _color_procedure(raw: str, label: str) -> str:
    """Color-code procedure: complex procedures stand out."""
    if raw in ("Competitive negotiated", "Competitive dialogue", "Innovation partnership"):
        return _yellow(label)
    if raw in ("Negotiated without publication", "Direct award"):
        return _style(label, "33;2")   # dim yellow — unntak
    return label                       # åpen/begrenset = normal


def _strip_html(text: str) -> str:
    """Strip HTML tags and normalize whitespace from API text fields."""
    # Replace block elements with newlines, then strip all tags
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"</(?:p|div|li|tr|h[1-6])>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"<[^>]+>", "", text)
    text = html.unescape(text)
    # Collapse whitespace: normalize runs of spaces/nbsp within lines
    lines = text.splitlines()
    lines = [re.sub(r"[ \t\xa0]+", " ", line).strip() for line in lines]
    # Collapse multiple blank lines
    result = []
    for line in lines:
        if line or (result and result[-1]):
            result.append(line)
    return "\n".join(result).strip()


def _truncate(text: str, width: int) -> str:
    return text if len(text) <= width else text[: width - 1] + "\u2026"


_COL_NR = 3
_COL_ID = 6
_COL_PROC = 12
_COL_THRESH = 14
_COL_FRIST = 12


def _print_procurement_list(procs: list[dict]) -> None:
    """Print a formatted, numbered list of procurements."""
    try:
        term_w = os.get_terminal_size().columns
    except OSError:
        term_w = 100
    fixed = 2 + _COL_NR + 2 + _COL_ID + 2 + _COL_PROC + _COL_THRESH + _COL_FRIST
    name_w = max(20, term_w - fixed)

    # Build header with padding BEFORE applying dim
    header = (
        f"{'#':>{_COL_NR}}  "
        f"{'ID':>{_COL_ID}}  "
        f"{'Prosedyre':<{_COL_PROC}}"
        f"{'Terskel':<{_COL_THRESH}}"
        f"{'Frist':<{_COL_FRIST}}"
        f"Navn"
    )
    print(f"\n  {_dim(header)}", file=sys.stderr)
    print(f"  {_dim('─' * min(term_w - 4, 96))}", file=sys.stderr)

    for i, p in enumerate(procs, 1):
        pid = p.get("id", "?")
        proc_raw = p.get("procedure", "")
        proc_label = PROCEDURE_SHORT.get(proc_raw, "?")
        thresh_raw = p.get("threshold") or ""
        thresh_label = THRESHOLD_SHORT.get(thresh_raw, thresh_raw or "?")
        deadline = _fmt_date(_get_timeline_date(p, "submission"))[:10]
        name = p.get("name") or "?"
        seq = p.get("sequenceId") or ""
        label = _truncate(name, name_w)
        seq_str = f" {_dim(seq)}" if seq else ""

        # Pad plain text first, then apply color so widths stay correct
        proc_padded = f"{proc_label:<{_COL_PROC}}"
        thresh_padded = f"{thresh_label:<{_COL_THRESH}}"
        frist_padded = f"{deadline:<{_COL_FRIST}}"

        print(
            f"  {_bold(f'{i:>{_COL_NR}}')}"
            f"  {_dim(f'{pid:>{_COL_ID}}')}"
            f"  {_color_procedure(proc_raw, proc_padded)}"
            f"{_color_threshold(thresh_raw, thresh_padded)}"
            f"{_dim(frist_padded)}"
            f"{label}{seq_str}",
            file=sys.stderr,
        )

    print(file=sys.stderr)


def _select_procurement(procs: list[dict]) -> dict:
    """Let user pick a procurement interactively."""
    _print_procurement_list(procs)
    while True:
        try:
            choice = input(f"  Velg anskaffelse {_dim(f'(1-{len(procs)})')}: ").strip()
        except (EOFError, KeyboardInterrupt):
            print(file=sys.stderr)
            sys.exit(0)
        if not choice:
            continue
        try:
            idx = int(choice) - 1
            if 0 <= idx < len(procs):
                return procs[idx]
        except ValueError:
            pass
        print(f"  {_yellow('?')} Skriv et tall mellom 1 og {len(procs)}", file=sys.stderr)


# -- Section generators ------------------------------------------------------

def _section_general_info(procurement: dict, activities: list[dict]) -> str:
    lines = []
    lines.append("## Generell informasjon")
    lines.append("")

    seq_id = procurement.get("sequenceId") or ""
    ext_id = procurement.get("externalId") or ""
    sak_ref = seq_id
    if ext_id:
        sak_ref += f" (ekstern ref: {ext_id})"

    procurer = procurement.get("about_procurer") or {}
    procurer_name = procurer.get("name") or ""
    national_id = procurer.get("national_id") or ""
    procurer_str = procurer_name
    if national_id:
        procurer_str += f" (org.nr. {national_id})"

    contact = procurer.get("contact_person") or ""

    name = _strip_html(procurement.get("name") or "")
    description = _strip_html(procurement.get("description") or "")
    # Collapse to single line for markdown table cell
    description = re.sub(r"\s*\n\s*", " ", description)
    desc_str = name
    if description and description.lower() != name.lower():
        desc_str += f". {description}"

    estimated_value = procurement.get("estimated_value")
    currency = procurement.get("currency") or "NOK"
    value_str = _fmt_currency(estimated_value, currency)

    duration_months = procurement.get("duration_months")
    duration = procurement.get("duration")
    if duration_months:
        duration_str = f"{duration_months} måneder"
    elif duration:
        duration_str = str(duration)
    else:
        duration_str = "<!-- MANUELT -->"

    submission_date = _get_timeline_date(procurement, "submission")
    submission_str = _fmt_datetime(submission_date) if submission_date else "<!-- MANUELT -->"

    lines.append("| Felt | Beskrivelse |")
    lines.append("| --- | --- |")
    lines.append(f"| **Konkurransens saksnummer:** | {sak_ref} |")
    lines.append(f"| **Oppdragsgiver(e):** | {procurer_str} |")
    lines.append(f"| **Protokollfører/saksbehandler:** | {contact} <!-- MANUELT: bekreft --> |")
    lines.append(f"| **Beskrivelse av anskaffelsen:** | {desc_str} |")
    lines.append(f"| **Anskaffelsens anslåtte verdi på kunngjøringstidspunktet:** | {value_str} |")
    lines.append(f"| **Kontraktens varighet:** | {duration_str} |")
    lines.append(f"| **Frist for innlevering av tilbud:** | {submission_str} |")
    lines.append("")

    submissions = _get_activities_by_action(activities, "SUBMIT_BID")
    lines.append("**Tidspunkt for mottak av tilbud:**")
    lines.append("")
    lines.append("| Leverandørens navn | Tidspunkt for mottak |")
    lines.append("| --- | --- |")
    if submissions:
        for s in submissions:
            org = s.get("organization") or {}
            org_name = org.get("name") or "Ukjent leverandør"
            date = _fmt_datetime(s.get("date"))
            lines.append(f"| {org_name} | {date} |")
    else:
        lines.append("| *Ingen tilbud registrert* | |")
    lines.append("")

    return "\n".join(lines)


def _section_procedure(procurement: dict, activities: list[dict]) -> str:
    lines = []
    lines.append("## Anskaffelsesprosedyre")
    lines.append("")
    lines.append("Følgende anskaffelsesprosedyre er lagt til grunn i denne konkurransen:")
    lines.append("")

    procedure = procurement.get("procedure") or ""
    selected = PROCEDURE_MAP.get(procedure, "")

    for p in ALL_PROCEDURES:
        if p == selected:
            lines.append(f"- **[x] {p}**")
        else:
            lines.append(f"- ☐ {p}")
    lines.append("")

    if procedure in ("Competitive negotiated", "Competitive dialogue"):
        lines.append("**Begrunnelse for bruk av konkurransepreget dialog eller konkurranse med forhandling etter forutgående kunngjøring:**")
        lines.append("<!-- MANUELT: Fyll inn begrunnelse for valg av prosedyre -->")
    else:
        lines.append("**Begrunnelse for bruk av konkurransepreget dialog eller konkurranse med forhandling etter forutgående kunngjøring:**")
        lines.append("Ikke relevant.")
    lines.append("")

    if procedure in ("Negotiated without publication", "Direct award"):
        code = procurement.get("direct_award_justification_code") or ""
        reason = procurement.get("direct_award_justification_reason") or "<!-- MANUELT -->"
        lines.append("**Begrunnelse for å bruke konkurranse med forhandling uten forutgående kunngjøring eller anskaffelse uten konkurranse:**")
        if code:
            lines.append(f"Hjemmel: {code}. {reason}")
        else:
            lines.append(reason)
    else:
        lines.append("**Begrunnelse for å bruke konkurranse med forhandling uten forutgående kunngjøring eller anskaffelse uten konkurranse:**")
        lines.append("Ikke relevant.")
    lines.append("")

    lines.append("**Dersom anskaffelsen ikke deles opp i delkontrakter, begrunnelse for ikke å dele opp kontrakten (jf. FOA § 19-4):**")
    lines.append("<!-- MANUELT: Fyll inn begrunnelse for hvorfor det ikke er delt opp i delkontrakter -->")
    lines.append("")

    doffin_activities = _get_activities_by_action(activities, "DOFFIN_NOTICE_STATUS_PUBLISHED")
    publish_activities = _get_activities_by_action(activities, "PUBLISH_TO_DOFFIN")

    announcement_date = ""
    doffin_ref = ""
    ted_ref = ""

    if publish_activities:
        announcement_date = _fmt_date(publish_activities[0].get("date"))

    if doffin_activities:
        desc = doffin_activities[0].get("description") or {}
        doffin_notice = desc.get("doffinNotice") or {}
        doffin_ref = doffin_notice.get("ngoj") or ""
        ted_ref = doffin_notice.get("publicationId") or ""
        if not announcement_date:
            announcement_date = _fmt_date(
                doffin_notice.get("publicationDate") or doffin_activities[0].get("date")
            )

    lines.append("**Kunngjøring:**")
    if announcement_date:
        parts = [f"Anskaffelsen ble kunngjort {announcement_date} på Doffin"]
        if doffin_ref:
            parts.append(f"referansenummer {doffin_ref} (NGOJ)")
        if ted_ref:
            parts.append(f"TED-referanse {ted_ref}")
        lines.append(", ".join(parts) + ".")
    else:
        pub_date = procurement.get("publicationDate")
        if pub_date:
            lines.append(f"Anskaffelsen ble kunngjort {_fmt_date(pub_date)}.")
        else:
            lines.append("<!-- MANUELT: Kunngjøringsinformasjon mangler -->")
    lines.append("")

    return "\n".join(lines)


def _section_formal_rejection(activities: list[dict]) -> str:
    lines = []
    lines.append("## Avvisning på grunn av formalfeil, jf. FOA § 24-1")
    lines.append("")

    rejections = _get_activities_by_action(activities, "REJECT_PARTICIPATION")
    if not rejections:
        lines.append("☐ Ingen leverandører eller tilbud ble avvist")
        lines.append("<!-- MANUELT: Bekreft at ingen ble avvist på formalfeil -->")
    else:
        lines.append("<!-- MANUELT: Avgjør hvilke avvisninger som gjelder formalfeil (§ 24-1) vs. kvalifikasjon (§ 24-2) -->")

    lines.append("")
    lines.append("| Leverandørens navn | Begrunnelsen for avvisningen | Dato sendt |")
    lines.append("| --- | --- | --- |")
    if rejections:
        for r in rejections:
            org = r.get("organization") or {}
            org_name = org.get("name") or "Ukjent"
            date = _fmt_date(r.get("date"))
            lines.append(f"| {org_name} | <!-- MANUELT: begrunnelse --> | {date} |")
    else:
        lines.append("| | | |")
    lines.append("")

    return "\n".join(lines)


def _section_preliminary_qualification(procedure: str) -> str:
    lines = []
    lines.append("## Foreløpig kvalifikasjonsvurdering, jf. FOA § 17-1 annet ledd")
    lines.append("")

    if procedure == "Open":
        lines.append("Ikke relevant (åpen anbudskonkurranse).")
    else:
        lines.append("**Leverandører som er kvalifisert:**")
        lines.append("<!-- MANUELT -->")
    lines.append("")

    return "\n".join(lines)


def _section_qualification() -> str:
    lines = []
    lines.append("## Kvalifikasjonsvurdering")
    lines.append("")
    lines.append("<!-- MANUELT: Kvalifikasjonskrav og -vurdering er ikke tilgjengelig via API. Fyll inn basert på konkurransegrunnlaget. -->")
    lines.append("")
    lines.append("**Leverandører som er kvalifisert:**")
    lines.append("<!-- MANUELT -->")
    lines.append("")
    lines.append("**Hvis relevant, begrunnelse for hvorfor leverandører som har restanser i henhold til skatte- og avgiftslovgivningen har fått delta i konkurransen:**")
    lines.append("<!-- MANUELT -->")
    lines.append("")

    return "\n".join(lines)


def _section_supplier_rejection(activities: list[dict]) -> str:
    lines = []
    lines.append("## Leverandører som er avvist, jf. FOA § 24-2")
    lines.append("")

    rejections = _get_activities_by_action(activities, "REJECT_PARTICIPATION")
    if not rejections:
        lines.append("☐ Ingen leverandører ble avvist")
        lines.append("<!-- MANUELT: Bekreft. API viser ingen avvisningshendelser. -->")
    else:
        lines.append("Følgende leverandører ble avvist:")

    lines.append("")
    lines.append("| Leverandørens navn | Begrunnelsen for avvisningen | Dato sendt |")
    lines.append("| --- | --- | --- |")
    if rejections:
        for r in rejections:
            org = r.get("organization") or {}
            org_name = org.get("name") or "Ukjent"
            date = _fmt_date(r.get("date"))
            lines.append(f"| {org_name} | <!-- MANUELT: begrunnelse --> | {date} |")
    else:
        lines.append("| | | |")
    lines.append("")

    return "\n".join(lines)


def _section_supplier_selection(procedure: str, activities: list[dict]) -> str:
    lines = []
    lines.append("## Utvelgelse av leverandører")
    lines.append("")

    if procedure == "Open":
        lines.append("Ikke relevant (åpen anbudskonkurranse — ingen utvelgelsesfase).")
    elif procedure in ("Limited", "Competitive negotiated", "Innovation partnership", "Competitive dialogue"):
        qualifying = _get_activities_by_action(activities, "QUALIFYING_PARTICIPANTS")
        if qualifying:
            lines.append("<!-- MANUELT: Fyll inn begrunnelse for utvelgelse per leverandør -->")
            lines.append("")
            lines.append("| Leverandørens navn | Begrunnelse for utvelgelse |")
            lines.append("| --- | --- |")
            for q in qualifying:
                desc = q.get("description") or {}
                tenders_ids = desc.get("tendersIds") or []
                for tid in tenders_ids:
                    lines.append(f"| Leverandør (tender {tid}) | <!-- MANUELT --> |")
        else:
            lines.append("<!-- MANUELT: Fyll inn utvelgelsesresultat -->")
            lines.append("")
            lines.append("| Leverandørens navn | Begrunnelse for utvelgelse |")
            lines.append("| --- | --- |")
            lines.append("| <!-- MANUELT --> | <!-- MANUELT --> |")
    else:
        lines.append("Ikke relevant.")
    lines.append("")

    return "\n".join(lines)


def _section_bid_rejection() -> str:
    lines = []
    lines.append("## Tilbud som er avvist, jf. FOA §§ 24-8 og 24-9")
    lines.append("")
    lines.append("☐ Ingen tilbud ble avvist")
    lines.append("<!-- MANUELT: Bekreft -->")
    lines.append("")
    lines.append("| Leverandørenes navn | Begrunnelsen for avvisningen | Dato sendt |")
    lines.append("| --- | --- | --- |")
    lines.append("| | | |")
    lines.append("")

    return "\n".join(lines)


def _section_clarification(procurement: dict, activities: list[dict]) -> str:
    lines = []
    lines.append("## Ettersending og avklaring av opplysninger og dokumentasjon, jf. FOA § 23-5")
    lines.append("")

    submission_deadline = _parse_submission_deadline(procurement)
    conversations = _get_activities_by_action(activities, "CONVERSATION_MARKED_COMPLETED")

    post_deadline_convs = []
    if submission_deadline and conversations:
        for c in conversations:
            conv_date_str = c.get("date")
            if conv_date_str:
                try:
                    conv_date = datetime.fromisoformat(conv_date_str.replace("Z", "+00:00"))
                    if conv_date > submission_deadline:
                        post_deadline_convs.append(c)
                except (ValueError, TypeError):
                    pass

    if not post_deadline_convs:
        lines.append("☐ Det ble ikke foretatt avklaringer eller dialog")
        if not conversations:
            lines.append("<!-- MANUELT: Bekreft. API har ingen avklaringshendelser for denne anskaffelsen. -->")
        else:
            lines.append("<!-- MANUELT: Bekreft. API har meldingshendelser, men alle er før tilbudsfrist (Q&A). -->")
    else:
        lines.append("Følgende avklaringer/ettersendinger ble gjennomført etter tilbudsfrist:")

    lines.append("")
    lines.append("| Leverandørens navn | Dato | Hvordan? |")
    lines.append("| --- | --- | --- |")

    if post_deadline_convs:
        for c in post_deadline_convs:
            org = c.get("organization") or {}
            org_name = org.get("name") or "Ukjent"
            date = _fmt_date(c.get("date"))
            desc = c.get("description") or {}
            title = desc.get("conversationTitle") or ""
            how = f"Melding i KGV: «{title}»" if title else "Melding i KGV"
            lines.append(f"| {org_name} | {date} | {how} |")
    else:
        lines.append("| | | |")
    lines.append("")

    return "\n".join(lines)


def _section_negotiations(procedure: str) -> str:
    lines = []
    lines.append("## Forhandlinger")
    lines.append("")

    if procedure in ("Competitive negotiated", "Innovation partnership"):
        lines.append("<!-- MANUELT: Fyll inn forhandlingsdetaljer -->")
        lines.append("")
        lines.append("☐ Det ble ikke gjennomført forhandlinger")
        lines.append("")
        lines.append("| Leverandørens navn | Dato for forhandling | Mottatt revidert tilbud |")
        lines.append("| --- | --- | --- |")
        lines.append("| <!-- MANUELT --> | <!-- MANUELT --> | <!-- MANUELT --> |")
    else:
        lines.append(f"Ikke relevant ({PROCEDURE_MAP.get(procedure, procedure)})." if procedure != "Open" else "Ikke relevant (åpen anbudskonkurranse).")
    lines.append("")

    return "\n".join(lines)


def _section_dialog(procedure: str) -> str:
    lines = []
    lines.append("## Dialog")
    lines.append("")

    if procedure == "Competitive dialogue":
        lines.append("<!-- MANUELT: Fyll inn dialogdetaljer -->")
        lines.append("")
        lines.append("☐ Det ble ikke gjennomført dialog")
        lines.append("")
        lines.append("| Leverandørens navn | Dato for dialog |")
        lines.append("| --- | --- |")
        lines.append("| <!-- MANUELT --> | <!-- MANUELT --> |")
    else:
        lines.append(f"Ikke relevant ({PROCEDURE_MAP.get(procedure, procedure)})." if procedure != "Open" else "Ikke relevant (åpen anbudskonkurranse).")
    lines.append("")

    return "\n".join(lines)


def _section_bids_in_evaluation(activities: list[dict]) -> str:
    lines = []
    lines.append("## Tilbud som er med i tildelingsvurderingen")
    lines.append("")

    submissions = _get_activities_by_action(activities, "SUBMIT_BID")
    rejections = _get_activities_by_action(activities, "REJECT_PARTICIPATION")
    withdrawals = _get_activities_by_action(activities, "WITHDRAW_PARTICIPATION")

    rejected_names = {
        (r.get("organization") or {}).get("name", "").lower()
        for r in rejections
    }
    withdrawn_names = {
        (w.get("organization") or {}).get("name", "").lower()
        for w in withdrawals
    }
    excluded = rejected_names | withdrawn_names

    evaluated = []
    for s in submissions:
        org = s.get("organization") or {}
        name = org.get("name") or "Ukjent"
        if name.lower() not in excluded:
            evaluated.append(name)

    lines.append("| Tilbuds-/løpenummer | Leverandørenes navn |")
    lines.append("| --- | --- |")
    if evaluated:
        for i, name in enumerate(evaluated, 1):
            lines.append(f"| {i} | {name} |")
    else:
        lines.append("| | <!-- MANUELT --> |")
    lines.append("")

    return "\n".join(lines)


def _section_award(procurement: dict, activities: list[dict]) -> str:
    lines = []
    lines.append("## Det (de) valgte tilbud med begrunnelse og kontraktsverdi")
    lines.append("")
    lines.append("<!-- MANUELT: Fyll inn navn på valgt leverandør, tildelingsbegrunnelse og kontraktsverdi. API gir kun tenderIds, ikke leverandørnavn eller begrunnelse. -->")
    lines.append("")

    total_value = procurement.get("contracts_total_value_amount")
    estimated = procurement.get("estimated_value")
    currency = procurement.get("currency") or "NOK"
    if total_value:
        lines.append(f"**Kontraktsverdi:** {_fmt_currency(total_value, currency)}")
    elif estimated:
        lines.append(f"**Kontraktsverdi:** {_fmt_currency(estimated, currency)} (estimert verdi)")
    else:
        lines.append("**Kontraktsverdi:** <!-- MANUELT -->")
    lines.append("")

    award_letters = procurement.get("areAwardLettersSent")

    lines.append("| Felt | Beskrivelse |")
    lines.append("| --- | --- |")
    if award_letters:
        lines.append("| **Meddelelsesbrev sendt:** | <!-- MANUELT: areAwardLettersSent=True, men dato ikke i API --> |")
    else:
        lines.append("| **Meddelelsesbrev sendt:** | <!-- MANUELT --> |")
    lines.append("| **Karensperiodens utløp:** | <!-- MANUELT --> |")
    lines.append("| **Eventuelle klager:** | <!-- MANUELT --> |")
    lines.append("| **Resultat av klage:** | <!-- MANUELT --> |")
    lines.append("")

    award_date = _get_timeline_date(procurement, "award decision")
    if award_date:
        lines.append(f"**Tildelingsbeslutning:** {_fmt_date(award_date)}")
    else:
        award_activities = _get_activities_by_action(activities, "AWARDING_PARTICIPANTS")
        if award_activities:
            lines.append(f"**Tildelingsbeslutning:** {_fmt_date(award_activities[0].get('date'))}")
        else:
            lines.append("**Tildelingsbeslutning:** <!-- MANUELT -->")
    lines.append("")

    return "\n".join(lines)


def _section_framework_agreement(procurement: dict) -> str:
    lines = []
    lines.append("## Tildeling av rammeavtaler")
    lines.append("")

    is_framework = procurement.get("framework_agreement_involved")
    if not is_framework:
        lines.append("Ikke relevant (dette er ikke en rammeavtale).")
    else:
        max_participants = procurement.get("framework_agreement_maximum_participants")
        if max_participants and int(max_participants) == 1:
            lines.append("- **Rammeavtale med en leverandør?** Ja")
            lines.append("- **Rammeavtale med flere leverandører?** Nei")
        elif max_participants and int(max_participants) > 1:
            lines.append("- **Rammeavtale med en leverandør?** Nei")
            lines.append(f"- **Rammeavtale med flere leverandører?** Ja (maks {max_participants} deltakere)")
            lines.append("- Ved rammeavtale med flere leverandører; Hvilken fordelingsmekanisme skal benyttes? <!-- MANUELT -->")
            lines.append("- Ved minikonkurranse som fordelingsmekanisme; Hvilke kriterier skal brukes? <!-- MANUELT -->")
        else:
            lines.append("- **Rammeavtale med en leverandør?** <!-- MANUELT -->")
            lines.append("- **Rammeavtale med flere leverandører?** <!-- MANUELT -->")
    lines.append("")

    return "\n".join(lines)


def _section_other(procurement: dict) -> str:
    lines = []
    lines.append("## Andre opplysninger og avslutning")
    lines.append("")
    lines.append("**Hvis relevant, hvilke deler av kontrakten valgte leverandør planlegger at underleverandører skal utføre, og underleverandørens navn, forutsatt at opplysningene er kjent:**")
    lines.append("<!-- MANUELT -->")
    lines.append("")

    is_cancelled = procurement.get("isCancelled")
    if is_cancelled:
        reason = procurement.get("cancelingReason") or "<!-- MANUELT: begrunnelse mangler -->"
        lines.append("**Hvis relevant, begrunnelse for hvorfor konkurransen avlyses, jf. FOA § 25-4:**")
        lines.append(reason)
    else:
        lines.append("**Hvis relevant, begrunnelse for hvorfor konkurransen avlyses, jf. FOA § 25-4:**")
        lines.append("Ikke relevant (konkurransen ble ikke avlyst).")
    lines.append("")

    lines.append("**Hvis relevant, opplysninger om tilfeller av inhabilitet eller konkurransevridning som følge av dialog med leverandørene, og eventuelle avhjelpende tiltak som er gjennomført:**")
    lines.append("<!-- MANUELT -->")
    lines.append("")

    lines.append("**Andre opplysninger, vesentlige forhold eller viktige beslutninger som er av betydning for konkurransen:**")
    lines.append("<!-- MANUELT -->")
    lines.append("")

    return "\n".join(lines)


def _section_data_quality(procurement: dict, activities: list[dict]) -> str:
    lines = []
    lines.append("## Datakvalitet — API vs. manuelt")
    lines.append("")

    submissions = _get_activities_by_action(activities, "SUBMIT_BID")
    rejections = _get_activities_by_action(activities, "REJECT_PARTICIPATION")
    doffin = _get_activities_by_action(activities, "DOFFIN_NOTICE_STATUS_PUBLISHED")
    publish = _get_activities_by_action(activities, "PUBLISH_TO_DOFFIN")

    submission_deadline = _parse_submission_deadline(procurement)
    conversations = _get_activities_by_action(activities, "CONVERSATION_MARKED_COMPLETED")
    post_deadline = []
    if submission_deadline:
        for c in conversations:
            try:
                dt = datetime.fromisoformat((c.get("date") or "").replace("Z", "+00:00"))
                if dt > submission_deadline:
                    post_deadline.append(c)
            except (ValueError, TypeError):
                pass

    lines.append("| Seksjon | Kilde | Merknad |")
    lines.append("| --- | --- | --- |")
    lines.append("| Generell informasjon | API | Komplett |")
    lines.append(f"| Tidspunkt for mottak | API (SUBMIT_BID) | {'Komplett' if submissions else 'Ingen tilbud registrert'} |")
    lines.append("| Prosedyretype | API (procedure) | Komplett |")
    lines.append(f"| Kunngjøring | API {'(DOFFIN_NOTICE)' if doffin or publish else ''} | {'Komplett' if doffin or publish else 'Trenger bekreftelse'} |")

    if rejections:
        lines.append(f"| Avvisning (formalfeil) | API ({len(rejections)} hendelser) | Trenger manuell klassifisering |")
    else:
        lines.append("| Avvisning (formalfeil) | API (ingen hendelser) | Trenger bekreftelse |")

    lines.append("| Kvalifikasjonsvurdering | **Manuelt** | Ikke i API |")

    if rejections:
        lines.append(f"| Avvisning av leverandører | API ({len(rejections)} hendelser) | Begrunnelse mangler |")
    else:
        lines.append("| Avvisning av leverandører | API (ingen hendelser) | Trenger bekreftelse |")

    lines.append("| Avvisning av tilbud | API (ingen hendelser) | Trenger bekreftelse |")

    if post_deadline:
        lines.append(f"| Ettersending/avklaring | API ({len(post_deadline)} meldinger etter frist) | Innhold mangler |")
    else:
        lines.append("| Ettersending/avklaring | API (ingen hendelser) | Trenger bekreftelse |")

    lines.append(f"| Tilbud i vurdering | API (SUBMIT_BID) | {'Komplett' if submissions else 'Ingen'} |")
    lines.append("| Valgte tilbud + begrunnelse | **Manuelt** | API har kun tenderIds |")
    lines.append("| Meddelelsesbrev/karens | **Manuelt** | Kun flag, ikke datoer |")
    lines.append("| Delkontrakter (begrunnelse) | **Manuelt** | Ikke i API |")
    lines.append("| Underleverandører | **Manuelt** | Ikke i API |")
    lines.append("| Inhabilitet | **Manuelt** | Ikke i API |")

    return "\n".join(lines)


# -- Docx helpers --------------------------------------------------------------

def _docx_setup(doc):
    """Configure document margins, styles, and page footer."""
    from docx.shared import Cm

    # Narrower margins (2.0 cm)
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Heading 2: add space before for section separation
    h2_style = doc.styles["Heading 2"]
    h2_style.paragraph_format.space_before = Pt(18)
    h2_style.paragraph_format.space_after = Pt(6)

    # Page numbers in footer
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = 1  # CENTER
        run = p.add_run()
        # PAGE field
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), r" PAGE ")
        run._element.append(fld)
        run2 = p.add_run(" / ")
        run2.font.size = Pt(8)
        run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run3 = p.add_run()
        fld2 = OxmlElement("w:fldSimple")
        fld2.set(qn("w:instr"), r" NUMPAGES ")
        run3._element.append(fld2)
        for r in [run, run3]:
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def _docx_shade_header_row(table):
    """Apply light grey background to the first row of a table."""
    for cell in table.rows[0].cells:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), "F2F2F2")
        cell._element.get_or_add_tcPr().append(shading)


def _docx_set_col_widths(table, widths_cm):
    """Set column widths in cm."""
    from docx.shared import Cm
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            if i < len(row.cells):
                row.cells[i].width = Cm(w)


def _add_manual(paragraph, text="[Fyll inn]"):
    """Add a yellow-highlighted, italic manual marker to a paragraph."""
    run = paragraph.add_run(text)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x66, 0x00)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "FFFF00")
    run._element.get_or_add_rPr().append(shading)
    return run


def _docx_add_table(doc, headers, rows):
    """Add a table with bold headers, grey header row, and bordered cells."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_data in enumerate(row_data):
            table.rows[r_idx + 1].cells[c_idx].text = str(cell_data)
    _docx_shade_header_row(table)
    return table


def _docx_add_table_with_manual(doc, headers, rows):
    """Add a table where cells containing None get a yellow manual marker."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_data in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if cell_data is None:
                cell.text = ""
                _add_manual(cell.paragraphs[0])
            else:
                cell.text = str(cell_data)
    _docx_shade_header_row(table)
    return table


def _docx_info_table(doc, info_rows):
    """Shared helper for the general info label-value table."""
    table = doc.add_table(rows=len(info_rows), cols=2)
    table.style = "Table Grid"
    _docx_set_col_widths(table, [6.0, 11.0])
    for i, (label, val) in enumerate(info_rows):
        cell_label = table.rows[i].cells[0]
        cell_label.text = ""
        run = cell_label.paragraphs[0].add_run(label)
        run.bold = True
        cell_val = table.rows[i].cells[1]
        if val is None:
            cell_val.text = ""
            _add_manual(cell_val.paragraphs[0])
        else:
            cell_val.text = str(val)
    return table


def _docx_general_info(doc, procurement, activities):
    doc.add_heading("Generell informasjon", level=2)

    seq_id = procurement.get("sequenceId") or ""
    ext_id = procurement.get("externalId") or ""
    sak_ref = seq_id
    if ext_id:
        sak_ref += f" (ekstern ref: {ext_id})"

    procurer = procurement.get("about_procurer") or {}
    procurer_name = procurer.get("name") or ""
    national_id = procurer.get("national_id") or ""
    procurer_str = procurer_name
    if national_id:
        procurer_str += f" (org.nr. {national_id})"

    contact = procurer.get("contact_person") or ""

    name = _strip_html(procurement.get("name") or "")
    description = _strip_html(procurement.get("description") or "")
    description = re.sub(r"\s*\n\s*", " ", description)
    desc_str = name
    if description and description.lower() != name.lower():
        desc_str += f". {description}"

    estimated_value = procurement.get("estimated_value")
    currency = procurement.get("currency") or "NOK"
    value_str = _fmt_currency(estimated_value, currency)

    duration_months = procurement.get("duration_months")
    duration = procurement.get("duration")
    if duration_months:
        duration_str = f"{duration_months} måneder"
    elif duration:
        duration_str = str(duration)
    else:
        duration_str = None

    submission_date = _get_timeline_date(procurement, "submission")
    submission_str = _fmt_datetime(submission_date) if submission_date else None

    _docx_info_table(doc, [
        ("Konkurransens saksnummer:", sak_ref),
        ("Oppdragsgiver(e):", procurer_str),
        ("Protokollfører/saksbehandler:", contact or None),
        ("Beskrivelse av anskaffelsen:", desc_str),
        ("Anskaffelsens anslåtte verdi på kunngjøringstidspunktet:", value_str),
        ("Kontraktens varighet:", duration_str),
        ("Frist for innlevering av tilbud:", submission_str),
    ])

    submissions = _get_activities_by_action(activities, "SUBMIT_BID")
    p = doc.add_paragraph()
    p.add_run("Tidspunkt for mottak av tilbud:").bold = True

    if submissions:
        rows = []
        for s in submissions:
            org = s.get("organization") or {}
            org_name = org.get("name") or "Ukjent leverandør"
            date = _fmt_datetime(s.get("date"))
            rows.append((org_name, date))
        _docx_add_table(doc, ["Leverandørens navn", "Tidspunkt for mottak"], rows)
    else:
        _docx_add_table(doc, ["Leverandørens navn", "Tidspunkt for mottak"],
                         [("Ingen tilbud registrert", "")])


def _docx_procedure(doc, procurement, activities):
    doc.add_heading("Anskaffelsesprosedyre", level=2)
    doc.add_paragraph("Følgende anskaffelsesprosedyre er lagt til grunn i denne konkurransen:")

    procedure = procurement.get("procedure") or ""
    selected = PROCEDURE_MAP.get(procedure, "")
    for p_name in ALL_PROCEDURES:
        if p_name == selected:
            doc.add_paragraph(f"\u2612 {p_name}").runs[0].bold = True
        else:
            doc.add_paragraph(f"\u2610 {p_name}")

    p = doc.add_paragraph()
    p.add_run("Begrunnelse for bruk av konkurransepreget dialog eller konkurranse med forhandling etter forutgående kunngjøring:").bold = True
    if procedure in ("Competitive negotiated", "Competitive dialogue"):
        p2 = doc.add_paragraph()
        _add_manual(p2, "[Fyll inn begrunnelse for valg av prosedyre]")
    else:
        doc.add_paragraph("Ikke relevant.")

    p = doc.add_paragraph()
    p.add_run("Begrunnelse for å bruke konkurranse med forhandling uten forutgående kunngjøring eller anskaffelse uten konkurranse:").bold = True
    if procedure in ("Negotiated without publication", "Direct award"):
        code = procurement.get("direct_award_justification_code") or ""
        reason = procurement.get("direct_award_justification_reason") or ""
        if code or reason:
            doc.add_paragraph(f"Hjemmel: {code}. {reason}" if code else reason)
        else:
            p2 = doc.add_paragraph()
            _add_manual(p2)
    else:
        doc.add_paragraph("Ikke relevant.")

    p = doc.add_paragraph()
    p.add_run("Dersom anskaffelsen ikke deles opp i delkontrakter, begrunnelse for ikke å dele opp kontrakten (jf. FOA \u00a7 19-4):").bold = True
    p2 = doc.add_paragraph()
    _add_manual(p2, "[Fyll inn begrunnelse]")

    doffin_activities = _get_activities_by_action(activities, "DOFFIN_NOTICE_STATUS_PUBLISHED")
    publish_activities = _get_activities_by_action(activities, "PUBLISH_TO_DOFFIN")
    announcement_date = ""
    doffin_ref = ""
    ted_ref = ""
    if publish_activities:
        announcement_date = _fmt_date(publish_activities[0].get("date"))
    if doffin_activities:
        desc = doffin_activities[0].get("description") or {}
        doffin_notice = desc.get("doffinNotice") or {}
        doffin_ref = doffin_notice.get("ngoj") or ""
        ted_ref = doffin_notice.get("publicationId") or ""
        if not announcement_date:
            announcement_date = _fmt_date(
                doffin_notice.get("publicationDate") or doffin_activities[0].get("date")
            )

    p = doc.add_paragraph()
    p.add_run("Kunngjøring:").bold = True
    if announcement_date:
        parts = [f"Anskaffelsen ble kunngjort {announcement_date} på Doffin"]
        if doffin_ref:
            parts.append(f"referansenummer {doffin_ref} (NGOJ)")
        if ted_ref:
            parts.append(f"TED-referanse {ted_ref}")
        doc.add_paragraph(", ".join(parts) + ".")
    else:
        pub_date = procurement.get("publicationDate")
        if pub_date:
            doc.add_paragraph(f"Anskaffelsen ble kunngjort {_fmt_date(pub_date)}.")
        else:
            p2 = doc.add_paragraph()
            _add_manual(p2, "[Kunngjøringsinformasjon mangler]")


def _docx_formal_rejection(doc, activities):
    doc.add_heading("Avvisning på grunn av formalfeil, jf. FOA \u00a7 24-1", level=2)

    rejections = _get_activities_by_action(activities, "REJECT_PARTICIPATION")
    if not rejections:
        doc.add_paragraph("\u2610 Ingen leverandører eller tilbud ble avvist")
        p = doc.add_paragraph()
        _add_manual(p, "[Bekreft at ingen ble avvist på formalfeil]")
    else:
        p = doc.add_paragraph()
        _add_manual(p, "[Avgjør hvilke avvisninger som gjelder formalfeil (\u00a7 24-1) vs. kvalifikasjon (\u00a7 24-2)]")

    rows = []
    if rejections:
        for r in rejections:
            org = r.get("organization") or {}
            org_name = org.get("name") or "Ukjent"
            date = _fmt_date(r.get("date"))
            rows.append((org_name, None, date))
    else:
        rows.append(("", "", ""))
    _docx_add_table_with_manual(doc,
        ["Leverandørens navn", "Begrunnelsen for avvisningen", "Dato sendt"], rows)


def _docx_preliminary_qualification(doc, procedure):
    doc.add_heading("Foreløpig kvalifikasjonsvurdering, jf. FOA \u00a7 17-1 annet ledd", level=2)

    if procedure == "Open":
        doc.add_paragraph("Ikke relevant (åpen anbudskonkurranse).")
    else:
        p = doc.add_paragraph()
        p.add_run("Leverandører som er kvalifisert:").bold = True
        p2 = doc.add_paragraph()
        _add_manual(p2)


def _docx_qualification(doc):
    doc.add_heading("Kvalifikasjonsvurdering", level=2)
    p = doc.add_paragraph()
    _add_manual(p, "[Kvalifikasjonskrav og -vurdering er ikke tilgjengelig via API. Fyll inn basert på konkurransegrunnlaget.]")

    p = doc.add_paragraph()
    p.add_run("Leverandører som er kvalifisert:").bold = True
    p2 = doc.add_paragraph()
    _add_manual(p2)

    p = doc.add_paragraph()
    p.add_run("Hvis relevant, begrunnelse for hvorfor leverandører som har restanser i henhold til skatte- og avgiftslovgivningen har fått delta i konkurransen:").bold = True
    p2 = doc.add_paragraph()
    _add_manual(p2)


def _docx_supplier_rejection(doc, activities):
    doc.add_heading("Leverandører som er avvist, jf. FOA \u00a7 24-2", level=2)

    rejections = _get_activities_by_action(activities, "REJECT_PARTICIPATION")
    if not rejections:
        doc.add_paragraph("\u2610 Ingen leverandører ble avvist")
        p = doc.add_paragraph()
        _add_manual(p, "[Bekreft. API viser ingen avvisningshendelser.]")
    else:
        doc.add_paragraph("Følgende leverandører ble avvist:")

    rows = []
    if rejections:
        for r in rejections:
            org = r.get("organization") or {}
            org_name = org.get("name") or "Ukjent"
            date = _fmt_date(r.get("date"))
            rows.append((org_name, None, date))
    else:
        rows.append(("", "", ""))
    _docx_add_table_with_manual(doc,
        ["Leverandørens navn", "Begrunnelsen for avvisningen", "Dato sendt"], rows)


def _docx_supplier_selection(doc, procedure, activities):
    doc.add_heading("Utvelgelse av leverandører", level=2)

    if procedure == "Open":
        doc.add_paragraph("Ikke relevant (åpen anbudskonkurranse — ingen utvelgelsesfase).")
    elif procedure in ("Limited", "Competitive negotiated", "Innovation partnership", "Competitive dialogue"):
        qualifying = _get_activities_by_action(activities, "QUALIFYING_PARTICIPANTS")
        p = doc.add_paragraph()
        _add_manual(p, "[Fyll inn begrunnelse for utvelgelse per leverandør]")
        if qualifying:
            rows = []
            for q in qualifying:
                desc = q.get("description") or {}
                tenders_ids = desc.get("tendersIds") or []
                for tid in tenders_ids:
                    rows.append((f"Leverandør (tender {tid})", None))
            if rows:
                _docx_add_table_with_manual(doc,
                    ["Leverandørens navn", "Begrunnelse for utvelgelse"], rows)
        else:
            _docx_add_table_with_manual(doc,
                ["Leverandørens navn", "Begrunnelse for utvelgelse"], [(None, None)])
    else:
        doc.add_paragraph("Ikke relevant.")


def _docx_bid_rejection(doc):
    doc.add_heading("Tilbud som er avvist, jf. FOA \u00a7\u00a7 24-8 og 24-9", level=2)
    doc.add_paragraph("\u2610 Ingen tilbud ble avvist")
    p = doc.add_paragraph()
    _add_manual(p, "[Bekreft]")
    _docx_add_table(doc,
        ["Leverandørenes navn", "Begrunnelsen for avvisningen", "Dato sendt"],
        [("", "", "")])


def _docx_clarification(doc, procurement, activities):
    doc.add_heading("Ettersending og avklaring av opplysninger og dokumentasjon, jf. FOA \u00a7 23-5", level=2)

    submission_deadline = _parse_submission_deadline(procurement)
    conversations = _get_activities_by_action(activities, "CONVERSATION_MARKED_COMPLETED")

    post_deadline_convs = []
    if submission_deadline and conversations:
        for c in conversations:
            conv_date_str = c.get("date")
            if conv_date_str:
                try:
                    conv_date = datetime.fromisoformat(conv_date_str.replace("Z", "+00:00"))
                    if conv_date > submission_deadline:
                        post_deadline_convs.append(c)
                except (ValueError, TypeError):
                    pass

    if not post_deadline_convs:
        doc.add_paragraph("\u2610 Det ble ikke foretatt avklaringer eller dialog")
        if not conversations:
            p = doc.add_paragraph()
            _add_manual(p, "[Bekreft. API har ingen avklaringshendelser for denne anskaffelsen.]")
        else:
            p = doc.add_paragraph()
            _add_manual(p, "[Bekreft. API har meldingshendelser, men alle er før tilbudsfrist (Q&A).]")
    else:
        doc.add_paragraph("Følgende avklaringer/ettersendinger ble gjennomført etter tilbudsfrist:")

    rows = []
    if post_deadline_convs:
        for c in post_deadline_convs:
            org = c.get("organization") or {}
            org_name = org.get("name") or "Ukjent"
            date = _fmt_date(c.get("date"))
            desc = c.get("description") or {}
            title = desc.get("conversationTitle") or ""
            how = f"Melding i KGV: \u00ab{title}\u00bb" if title else "Melding i KGV"
            rows.append((org_name, date, how))
    else:
        rows.append(("", "", ""))
    _docx_add_table(doc, ["Leverandørens navn", "Dato", "Hvordan?"], rows)


def _docx_negotiations(doc, procedure):
    doc.add_heading("Forhandlinger", level=2)

    if procedure in ("Competitive negotiated", "Innovation partnership"):
        p = doc.add_paragraph()
        _add_manual(p, "[Fyll inn forhandlingsdetaljer]")
        doc.add_paragraph("\u2610 Det ble ikke gjennomført forhandlinger")
        _docx_add_table_with_manual(doc,
            ["Leverandørens navn", "Dato for forhandling", "Mottatt revidert tilbud"],
            [(None, None, None)])
    else:
        label = PROCEDURE_MAP.get(procedure, procedure) if procedure != "Open" else "åpen anbudskonkurranse"
        doc.add_paragraph(f"Ikke relevant ({label}).")


def _docx_dialog(doc, procedure):
    doc.add_heading("Dialog", level=2)

    if procedure == "Competitive dialogue":
        p = doc.add_paragraph()
        _add_manual(p, "[Fyll inn dialogdetaljer]")
        doc.add_paragraph("\u2610 Det ble ikke gjennomført dialog")
        _docx_add_table_with_manual(doc,
            ["Leverandørens navn", "Dato for dialog"],
            [(None, None)])
    else:
        label = PROCEDURE_MAP.get(procedure, procedure) if procedure != "Open" else "åpen anbudskonkurranse"
        doc.add_paragraph(f"Ikke relevant ({label}).")


def _docx_bids_in_evaluation(doc, activities):
    doc.add_heading("Tilbud som er med i tildelingsvurderingen", level=2)

    submissions = _get_activities_by_action(activities, "SUBMIT_BID")
    rejections = _get_activities_by_action(activities, "REJECT_PARTICIPATION")
    withdrawals = _get_activities_by_action(activities, "WITHDRAW_PARTICIPATION")

    rejected_names = {
        (r.get("organization") or {}).get("name", "").lower()
        for r in rejections
    }
    withdrawn_names = {
        (w.get("organization") or {}).get("name", "").lower()
        for w in withdrawals
    }
    excluded = rejected_names | withdrawn_names

    evaluated = []
    for s in submissions:
        org = s.get("organization") or {}
        name = org.get("name") or "Ukjent"
        if name.lower() not in excluded:
            evaluated.append(name)

    if evaluated:
        rows = [(str(i), name) for i, name in enumerate(evaluated, 1)]
    else:
        rows = [("", None)]
    _docx_add_table_with_manual(doc, ["Tilbuds-/løpenummer", "Leverandørenes navn"], rows)


def _docx_award(doc, procurement, activities):
    doc.add_heading("Det (de) valgte tilbud med begrunnelse og kontraktsverdi", level=2)
    p = doc.add_paragraph()
    _add_manual(p, "[Fyll inn navn på valgt leverandør, tildelingsbegrunnelse og kontraktsverdi. API gir kun tenderIds, ikke leverandørnavn eller begrunnelse.]")

    total_value = procurement.get("contracts_total_value_amount")
    estimated = procurement.get("estimated_value")
    currency = procurement.get("currency") or "NOK"
    p = doc.add_paragraph()
    p.add_run("Kontraktsverdi: ").bold = True
    if total_value:
        p.add_run(_fmt_currency(total_value, currency))
    elif estimated:
        p.add_run(f"{_fmt_currency(estimated, currency)} (estimert verdi)")
    else:
        _add_manual(p)

    award_letters = procurement.get("areAwardLettersSent")
    _docx_info_table(doc, [
        ("Meddelelsesbrev sendt:", "Sendt (dato ikke tilgjengelig i API)" if award_letters else None),
        ("Karensperiodens utløp:", None),
        ("Eventuelle klager:", None),
        ("Resultat av klage:", None),
    ])

    award_date = _get_timeline_date(procurement, "award decision")
    p = doc.add_paragraph()
    p.add_run("Tildelingsbeslutning: ").bold = True
    if award_date:
        p.add_run(_fmt_date(award_date))
    else:
        award_activities = _get_activities_by_action(activities, "AWARDING_PARTICIPANTS")
        if award_activities:
            p.add_run(_fmt_date(award_activities[0].get("date")))
        else:
            _add_manual(p)


def _docx_framework_agreement(doc, procurement):
    doc.add_heading("Tildeling av rammeavtaler", level=2)

    is_framework = procurement.get("framework_agreement_involved")
    if not is_framework:
        doc.add_paragraph("Ikke relevant (dette er ikke en rammeavtale).")
        return

    max_participants = procurement.get("framework_agreement_maximum_participants")
    if max_participants and int(max_participants) == 1:
        doc.add_paragraph("Rammeavtale med en leverandør? Ja")
        doc.add_paragraph("Rammeavtale med flere leverandører? Nei")
    elif max_participants and int(max_participants) > 1:
        doc.add_paragraph("Rammeavtale med en leverandør? Nei")
        doc.add_paragraph(f"Rammeavtale med flere leverandører? Ja (maks {max_participants} deltakere)")
        p = doc.add_paragraph()
        p.add_run("Ved rammeavtale med flere leverandører; Hvilken fordelingsmekanisme skal benyttes? ")
        _add_manual(p)
        p = doc.add_paragraph()
        p.add_run("Ved minikonkurranse som fordelingsmekanisme; Hvilke kriterier skal brukes? ")
        _add_manual(p)
    else:
        p1 = doc.add_paragraph("Rammeavtale med en leverandør? ")
        _add_manual(p1)
        p2 = doc.add_paragraph("Rammeavtale med flere leverandører? ")
        _add_manual(p2)


def _docx_other(doc, procurement):
    doc.add_heading("Andre opplysninger og avslutning", level=2)

    p = doc.add_paragraph()
    p.add_run("Hvis relevant, hvilke deler av kontrakten valgte leverandør planlegger at underleverandører skal utføre, og underleverandørens navn, forutsatt at opplysningene er kjent:").bold = True
    p2 = doc.add_paragraph()
    _add_manual(p2)

    p = doc.add_paragraph()
    p.add_run("Hvis relevant, begrunnelse for hvorfor konkurransen avlyses, jf. FOA \u00a7 25-4:").bold = True
    is_cancelled = procurement.get("isCancelled")
    if is_cancelled:
        reason = procurement.get("cancelingReason") or ""
        if reason:
            doc.add_paragraph(reason)
        else:
            p2 = doc.add_paragraph()
            _add_manual(p2, "[Begrunnelse mangler]")
    else:
        doc.add_paragraph("Ikke relevant (konkurransen ble ikke avlyst).")

    p = doc.add_paragraph()
    p.add_run("Hvis relevant, opplysninger om tilfeller av inhabilitet eller konkurransevridning som følge av dialog med leverandørene, og eventuelle avhjelpende tiltak som er gjennomført:").bold = True
    p2 = doc.add_paragraph()
    _add_manual(p2)

    p = doc.add_paragraph()
    p.add_run("Andre opplysninger, vesentlige forhold eller viktige beslutninger som er av betydning for konkurransen:").bold = True
    p2 = doc.add_paragraph()
    _add_manual(p2)


def _docx_data_quality(doc, procurement, activities):
    doc.add_heading("Datakvalitet — API vs. manuelt", level=2)

    submissions = _get_activities_by_action(activities, "SUBMIT_BID")
    rejections = _get_activities_by_action(activities, "REJECT_PARTICIPATION")
    doffin = _get_activities_by_action(activities, "DOFFIN_NOTICE_STATUS_PUBLISHED")
    publish = _get_activities_by_action(activities, "PUBLISH_TO_DOFFIN")

    submission_deadline = _parse_submission_deadline(procurement)
    conversations = _get_activities_by_action(activities, "CONVERSATION_MARKED_COMPLETED")
    post_deadline = []
    if submission_deadline:
        for c in conversations:
            try:
                dt = datetime.fromisoformat((c.get("date") or "").replace("Z", "+00:00"))
                if dt > submission_deadline:
                    post_deadline.append(c)
            except (ValueError, TypeError):
                pass

    rows = [
        ("Generell informasjon", "API", "Komplett"),
        ("Tidspunkt for mottak", "API (SUBMIT_BID)", "Komplett" if submissions else "Ingen tilbud registrert"),
        ("Prosedyretype", "API (procedure)", "Komplett"),
        ("Kunngjøring", f"API {'(DOFFIN_NOTICE)' if doffin or publish else ''}", "Komplett" if doffin or publish else "Trenger bekreftelse"),
        ("Avvisning (formalfeil)", f"API ({len(rejections)} hendelser)" if rejections else "API (ingen hendelser)", "Trenger manuell klassifisering" if rejections else "Trenger bekreftelse"),
        ("Kvalifikasjonsvurdering", "Manuelt", "Ikke i API"),
        ("Avvisning av leverandører", f"API ({len(rejections)} hendelser)" if rejections else "API (ingen hendelser)", "Begrunnelse mangler" if rejections else "Trenger bekreftelse"),
        ("Avvisning av tilbud", "API (ingen hendelser)", "Trenger bekreftelse"),
        ("Ettersending/avklaring", f"API ({len(post_deadline)} meldinger etter frist)" if post_deadline else "API (ingen hendelser)", "Innhold mangler" if post_deadline else "Trenger bekreftelse"),
        ("Tilbud i vurdering", "API (SUBMIT_BID)", "Komplett" if submissions else "Ingen"),
        ("Valgte tilbud + begrunnelse", "Manuelt", "API har kun tenderIds"),
        ("Meddelelsesbrev/karens", "Manuelt", "Kun flag, ikke datoer"),
        ("Delkontrakter (begrunnelse)", "Manuelt", "Ikke i API"),
        ("Underleverandører", "Manuelt", "Ikke i API"),
        ("Inhabilitet", "Manuelt", "Ikke i API"),
    ]
    _docx_add_table(doc, ["Seksjon", "Kilde", "Merknad"], rows)


# -- Del II docx sections (nasjonal, under EØS) --------------------------------

def _docx_del2_general_info(doc, procurement, activities):
    """Same as Del III but title says 'del II'."""
    doc.add_heading("Generell informasjon", level=2)

    seq_id = procurement.get("sequenceId") or ""
    ext_id = procurement.get("externalId") or ""
    sak_ref = seq_id
    if ext_id:
        sak_ref += f" (ekstern ref: {ext_id})"

    procurer = procurement.get("about_procurer") or {}
    procurer_name = procurer.get("name") or ""
    national_id = procurer.get("national_id") or ""
    procurer_str = procurer_name
    if national_id:
        procurer_str += f" (org.nr. {national_id})"

    contact = procurer.get("contact_person") or ""

    name = _strip_html(procurement.get("name") or "")
    description = _strip_html(procurement.get("description") or "")
    description = re.sub(r"\s*\n\s*", " ", description)
    desc_str = name
    if description and description.lower() != name.lower():
        desc_str += f". {description}"

    estimated_value = procurement.get("estimated_value")
    currency = procurement.get("currency") or "NOK"
    value_str = _fmt_currency(estimated_value, currency)

    duration_months = procurement.get("duration_months")
    duration = procurement.get("duration")
    if duration_months:
        duration_str = f"{duration_months} måneder"
    elif duration:
        duration_str = str(duration)
    else:
        duration_str = None

    submission_date = _get_timeline_date(procurement, "submission")
    submission_str = _fmt_datetime(submission_date) if submission_date else None

    _docx_info_table(doc, [
        ("Konkurransens saksnummer:", sak_ref),
        ("Oppdragsgiver(e):", procurer_str),
        ("Protokollfører/saksbehandler:", contact or None),
        ("Beskrivelse av anskaffelsen:", desc_str),
        ("Kontraktens anslåtte verdi på kunngjøringstidspunktet:", value_str),
        ("Kontraktens varighet:", duration_str),
        ("Frist for innlevering av tilbud:", submission_str),
    ])

    submissions = _get_activities_by_action(activities, "SUBMIT_BID")
    p = doc.add_paragraph()
    p.add_run("Tidspunkt for mottak av tilbud:").bold = True

    if submissions:
        rows = []
        for s in submissions:
            org = s.get("organization") or {}
            org_name = org.get("name") or "Ukjent leverandør"
            date = _fmt_datetime(s.get("date"))
            rows.append((org_name, date))
        _docx_add_table(doc, ["Leverandørens navn", "Tidspunkt for mottak"], rows)
    else:
        _docx_add_table(doc, ["Leverandørens navn", "Tidspunkt for mottak"],
                         [("Ingen tilbud registrert", "")])


def _docx_del2_procedure(doc, procurement, activities):
    """Del II: Only 2 procedure types, no delkontrakter."""
    doc.add_heading("Anskaffelsesprosedyre", level=2)
    doc.add_paragraph("Følgende anskaffelsesprosedyre er lagt til grunn i denne konkurransen:")

    procedure = procurement.get("procedure") or ""
    selected = DEL2_PROCEDURE_MAP.get(procedure, "")

    for p_name in ALL_DEL2_PROCEDURES:
        if p_name == selected:
            doc.add_paragraph(f"\u2612 {p_name}").runs[0].bold = True
        else:
            doc.add_paragraph(f"\u2610 {p_name}")

    # Kunngjøring
    doffin_activities = _get_activities_by_action(activities, "DOFFIN_NOTICE_STATUS_PUBLISHED")
    publish_activities = _get_activities_by_action(activities, "PUBLISH_TO_DOFFIN")
    announcement_date = ""
    doffin_ref = ""
    ted_ref = ""
    if publish_activities:
        announcement_date = _fmt_date(publish_activities[0].get("date"))
    if doffin_activities:
        desc = doffin_activities[0].get("description") or {}
        doffin_notice = desc.get("doffinNotice") or {}
        doffin_ref = doffin_notice.get("ngoj") or ""
        ted_ref = doffin_notice.get("publicationId") or ""
        if not announcement_date:
            announcement_date = _fmt_date(
                doffin_notice.get("publicationDate") or doffin_activities[0].get("date")
            )

    p = doc.add_paragraph()
    p.add_run("Kunngjøring:").bold = True
    if announcement_date:
        parts = [f"Anskaffelsen ble kunngjort {announcement_date} i DOFFIN"]
        if doffin_ref:
            parts.append(f"referansenummer {doffin_ref}")
        if ted_ref:
            parts.append(f"TED-referanse {ted_ref}")
        doc.add_paragraph(", ".join(parts) + ".")
    else:
        pub_date = procurement.get("publicationDate")
        if pub_date:
            doc.add_paragraph(f"Anskaffelsen ble kunngjort {_fmt_date(pub_date)}.")
        else:
            p2 = doc.add_paragraph()
            _add_manual(p2, "[Kunngjøringsinformasjon mangler]")


def _docx_del2_dialog(doc, procurement, activities):
    """Del II: Dialog og/eller forhandlinger, jf. FOA § 9-3."""
    doc.add_heading("Dialog og/eller forhandlinger, jf. FOA \u00a7 9-3", level=2)

    p = doc.add_paragraph()
    p.add_run("Hvis relevant, begrunnelse for å fravike opplysningene i anskaffelsesdokumentene om planlagt dialog:").bold = True
    p2 = doc.add_paragraph()
    _add_manual(p2)

    doc.add_paragraph("\u2610 Ingen dialog eller forhandlinger gjennomført")
    p = doc.add_paragraph()
    _add_manual(p, "[Bekreft]")

    _docx_add_table_with_manual(doc,
        ["Leverandørens navn", "Dato", "Begrunnelse for hvorfor leverandøren er valgt ut til dialog"],
        [(None, None, None)])

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Hvis relevant, opplysninger om tilfeller av inhabilitet eller konkurransevridning som følge av dialog med leverandørene, og eventuelle avhjelpende tiltak som er gjennomført:").bold = True
    p2 = doc.add_paragraph()
    _add_manual(p2)


def _docx_del2_formal_rejection(doc, activities):
    """Del II: Avvisning formalfeil, jf. FOA § 9-4."""
    doc.add_heading("Avvisning på grunn av formalfeil, jf. FOA \u00a7 9-4", level=2)

    rejections = _get_activities_by_action(activities, "REJECT_PARTICIPATION")
    if not rejections:
        doc.add_paragraph("\u2610 Ingen avvist på grunn av formalfeil")
        p = doc.add_paragraph()
        _add_manual(p, "[Bekreft at ingen ble avvist på formalfeil]")
    else:
        p = doc.add_paragraph()
        _add_manual(p, "[Avgjør hvilke avvisninger som gjelder formalfeil (\u00a7 9-4) vs. kvalifikasjon (\u00a7 9-5)]")

    rows = []
    if rejections:
        for r in rejections:
            org = r.get("organization") or {}
            org_name = org.get("name") or "Ukjent"
            date = _fmt_date(r.get("date"))
            rows.append((org_name, None, date))
    else:
        rows.append(("", "", ""))
    _docx_add_table_with_manual(doc,
        ["Leverandørens navn", "Begrunnelsen for avvisningen", "Begrunnelse ble sendt"], rows)


def _docx_del2_qualification(doc):
    """Del II: Three-tier qualification (full, preliminary § 8-10, chosen supplier)."""
    # Tier 1: Full qualification
    doc.add_heading("Kvalifikasjonsvurdering", level=2)
    doc.add_paragraph("Matrisen strykes hvis ikke relevant. Matrisen benyttes hvis det ikke kreves egenerklæring som et foreløpig dokumentasjonsbevis for at leverandøren oppfyller kvalifikasjonskravene.")
    p = doc.add_paragraph()
    _add_manual(p, "[Kvalifikasjonskrav og -vurdering. Fyll inn basert på konkurransegrunnlaget.]")

    p = doc.add_paragraph()
    p.add_run("Leverandører som er kvalifisert:").bold = True
    p2 = doc.add_paragraph()
    _add_manual(p2)

    p = doc.add_paragraph()
    p.add_run("Hvis relevant, begrunnelse for hvorfor leverandører som har restanser i henhold til skatte- og avgiftslovgivningen har fått delta i konkurransen:").bold = True
    p2 = doc.add_paragraph()
    _add_manual(p2)

    # Tier 2: Preliminary qualification via self-declaration
    doc.add_heading("Foreløpig kvalifikasjonsvurdering", level=2)
    doc.add_paragraph("Matrisen strykes hvis ikke relevant. Matrisen benyttes hvis det kreves egenerklæring som et foreløpig dokumentasjonsbevis for at leverandøren oppfyller kvalifikasjonskravene, jf. FOA \u00a7 8-10 (1).")

    p = doc.add_paragraph()
    p.add_run("Leverandører som er kvalifisert:").bold = True
    p2 = doc.add_paragraph()
    _add_manual(p2)

    # Tier 3: Qualification of chosen supplier(s)
    doc.add_heading("Kvalifikasjonsvurdering av valgte leverandør(er)", level=2)
    doc.add_paragraph("Matrisen strykes hvis ikke relevant. Matrisen benyttes hvis det kreves egenerklæring som et foreløpig dokumentasjonsbevis. Oppdragsgiver skal kreve at valgte leverandør leverer oppdaterte dokumentasjonsbevis, jf. FOA \u00a7 8-10 (2).")

    p = doc.add_paragraph()
    p.add_run("Leverandør(er) som er kvalifisert:").bold = True
    p2 = doc.add_paragraph()
    _add_manual(p2)

    p = doc.add_paragraph()
    p.add_run("Hvis relevant, begrunnelse for hvorfor leverandør som har restanser i henhold til skatte- og avgiftslovgivningen har fått delta i konkurransen:").bold = True
    p2 = doc.add_paragraph()
    _add_manual(p2)


def _docx_del2_supplier_rejection(doc, activities):
    """Del II: Leverandører avvist, jf. FOA § 9-5."""
    doc.add_heading("Leverandører som er avvist, jf. FOA \u00a7 9-5", level=2)

    rejections = _get_activities_by_action(activities, "REJECT_PARTICIPATION")
    if not rejections:
        doc.add_paragraph("\u2610 Ingen leverandører avvist")
        p = doc.add_paragraph()
        _add_manual(p, "[Bekreft. API viser ingen avvisningshendelser.]")
    else:
        doc.add_paragraph("Følgende leverandører ble avvist:")

    rows = []
    if rejections:
        for r in rejections:
            org = r.get("organization") or {}
            org_name = org.get("name") or "Ukjent"
            date = _fmt_date(r.get("date"))
            rows.append((org_name, None, date))
    else:
        rows.append(("", "", ""))
    _docx_add_table_with_manual(doc,
        ["Leverandørens navn", "Begrunnelsen for avvisningen", "Begrunnelse ble sendt"], rows)


def _docx_del2_supplier_selection(doc, procedure, activities):
    """Del II: Utvelgelse — only for begrenset tilbudskonkurranse."""
    doc.add_heading("Utvelgelse av leverandører", level=2)

    if procedure != "Limited":
        doc.add_paragraph("Matrisen strykes hvis ikke relevant. Gjelder kun ved begrenset tilbudskonkurranse.")
        doc.add_paragraph("Ikke relevant (åpen tilbudskonkurranse).")
        return

    doc.add_paragraph("Gjelder kun ved begrenset tilbudskonkurranse.")
    qualifying = _get_activities_by_action(activities, "QUALIFYING_PARTICIPANTS")
    p = doc.add_paragraph()
    _add_manual(p, "[Fyll inn begrunnelse for utvelgelse per leverandør]")
    if qualifying:
        rows = []
        for q in qualifying:
            desc = q.get("description") or {}
            tenders_ids = desc.get("tendersIds") or []
            for tid in tenders_ids:
                rows.append((f"Leverandør (tender {tid})", None))
        if rows:
            _docx_add_table_with_manual(doc,
                ["Leverandørens navn", "Begrunnelse for utvelgelse"], rows)
    else:
        _docx_add_table_with_manual(doc,
            ["Leverandørens navn", "Begrunnelse for utvelgelse"], [(None, None)])


def _docx_del2_bid_rejection(doc):
    """Del II: Tilbud avvist, jf. FOA § 9-6."""
    doc.add_heading("Tilbud som er avvist, jf. FOA \u00a7 9-6", level=2)
    doc.add_paragraph("\u2610 Ingen tilbud avvist")
    p = doc.add_paragraph()
    _add_manual(p, "[Bekreft]")
    _docx_add_table(doc,
        ["Leverandørens navn", "Begrunnelsen for avvisningen", "Begrunnelse ble sendt"],
        [("", "", "")])


def _docx_del2_clarification(doc, procurement, activities):
    """Del II: Ettersending og avklaring (basert på grunnleggende prinsipper)."""
    doc.add_heading("Ettersending og avklaring", level=2)

    submission_deadline = _parse_submission_deadline(procurement)
    conversations = _get_activities_by_action(activities, "CONVERSATION_MARKED_COMPLETED")

    post_deadline_convs = []
    if submission_deadline and conversations:
        for c in conversations:
            conv_date_str = c.get("date")
            if conv_date_str:
                try:
                    conv_date = datetime.fromisoformat(conv_date_str.replace("Z", "+00:00"))
                    if conv_date > submission_deadline:
                        post_deadline_convs.append(c)
                except (ValueError, TypeError):
                    pass

    if not post_deadline_convs:
        doc.add_paragraph("\u2610 Det ble ikke foretatt avklaringer eller ettersendinger")
        if not conversations:
            p = doc.add_paragraph()
            _add_manual(p, "[Bekreft. API har ingen avklaringshendelser for denne anskaffelsen.]")
        else:
            p = doc.add_paragraph()
            _add_manual(p, "[Bekreft. API har meldingshendelser, men alle er før tilbudsfrist (Q&A).]")
    else:
        doc.add_paragraph("Følgende avklaringer/ettersendinger ble gjennomført etter tilbudsfrist:")

    rows = []
    if post_deadline_convs:
        for c in post_deadline_convs:
            org = c.get("organization") or {}
            org_name = org.get("name") or "Ukjent"
            date = _fmt_date(c.get("date"))
            desc = c.get("description") or {}
            title = desc.get("conversationTitle") or ""
            how = f"Melding i KGV: \u00ab{title}\u00bb" if title else "Melding i KGV"
            rows.append((org_name, date, how))
    else:
        rows.append(("", "", ""))
    _docx_add_table(doc, ["Leverandørens navn", "Dato", "Hvordan?"], rows)


def _docx_del2_award(doc, procurement, activities):
    """Del II: Valgt tilbud med begrunnelse og kontraktsverdi."""
    doc.add_heading("Det (de) valgte tilbud med begrunnelse og kontraktsverdi", level=2)

    doc.add_paragraph("Begrunnelsen skal inneholde tilstrekkelig informasjon om det valgte tilbud til at leverandøren kan vurdere om oppdragsgivers valg har vært saklig og forsvarlig. Begrunnelsen skal inneholde opplysninger om navnet på valgte leverandør, valgte tilbuds verdi, relative fordeler og egenskaper i forhold til de øvrige tilbudene.")

    p = doc.add_paragraph()
    _add_manual(p, "[Fyll inn navn på valgt leverandør, tildelingsbegrunnelse og kontraktsverdi.]")

    total_value = procurement.get("contracts_total_value_amount")
    estimated = procurement.get("estimated_value")
    currency = procurement.get("currency") or "NOK"
    p = doc.add_paragraph()
    p.add_run("Kontraktsverdi: ").bold = True
    if total_value:
        p.add_run(_fmt_currency(total_value, currency))
    elif estimated:
        p.add_run(f"{_fmt_currency(estimated, currency)} (estimert verdi)")
    else:
        _add_manual(p)

    award_date = _get_timeline_date(procurement, "award decision")
    p = doc.add_paragraph()
    p.add_run("Tildelingsbeslutning: ").bold = True
    if award_date:
        p.add_run(_fmt_date(award_date))
    else:
        award_activities = _get_activities_by_action(activities, "AWARDING_PARTICIPANTS")
        if award_activities:
            p.add_run(_fmt_date(award_activities[0].get("date")))
        else:
            _add_manual(p)


def _docx_del2_award_notification(doc, procurement):
    """Del II: Meddelelse om tildeling og klagefrist (not karensperiode)."""
    doc.add_heading("Meddelelse om tildeling og klagefrist før inngåelse av kontrakt", level=2)

    award_letters = procurement.get("areAwardLettersSent")
    _docx_info_table(doc, [
        ("Meddelelsesbrev sendt:", "Sendt (dato ikke tilgjengelig i API)" if award_letters else None),
        ("Klagefrist:", None),
        ("Eventuelle klager:", None),
        ("Resultat av klage:", None),
    ])


def _docx_del2_other(doc, procurement):
    """Del II: Andre opplysninger — § 10-4 for avlysning."""
    doc.add_heading("Andre opplysninger og avslutning", level=2)

    p = doc.add_paragraph()
    p.add_run("Hvis relevant, hvilke deler av kontrakten valgte leverandør planlegger at underleverandører skal utføre, og underleverandørens navn, forutsatt at opplysningene er kjent:").bold = True
    p2 = doc.add_paragraph()
    _add_manual(p2)

    p = doc.add_paragraph()
    p.add_run("Hvis relevant, begrunnelse for hvorfor konkurransen avlyses, jf. FOA \u00a7 10-4 (1):").bold = True
    is_cancelled = procurement.get("isCancelled")
    if is_cancelled:
        reason = procurement.get("cancelingReason") or ""
        if reason:
            doc.add_paragraph(reason)
        else:
            p2 = doc.add_paragraph()
            _add_manual(p2, "[Begrunnelse mangler]")
    else:
        doc.add_paragraph("Ikke relevant (konkurransen ble ikke avlyst).")

    p = doc.add_paragraph()
    p.add_run("Andre opplysninger, vesentlige forhold eller viktige beslutninger som er av betydning for konkurransen:").bold = True
    p2 = doc.add_paragraph()
    _add_manual(p2)


def _docx_del2_data_quality(doc, procurement, activities):
    """Del II: Datakvalitet — adjusted for Del II sections."""
    doc.add_heading("Datakvalitet — API vs. manuelt", level=2)

    submissions = _get_activities_by_action(activities, "SUBMIT_BID")
    rejections = _get_activities_by_action(activities, "REJECT_PARTICIPATION")
    doffin = _get_activities_by_action(activities, "DOFFIN_NOTICE_STATUS_PUBLISHED")
    publish = _get_activities_by_action(activities, "PUBLISH_TO_DOFFIN")

    submission_deadline = _parse_submission_deadline(procurement)
    conversations = _get_activities_by_action(activities, "CONVERSATION_MARKED_COMPLETED")
    post_deadline = []
    if submission_deadline:
        for c in conversations:
            try:
                dt = datetime.fromisoformat((c.get("date") or "").replace("Z", "+00:00"))
                if dt > submission_deadline:
                    post_deadline.append(c)
            except (ValueError, TypeError):
                pass

    rows = [
        ("Generell informasjon", "API", "Komplett"),
        ("Tidspunkt for mottak", "API (SUBMIT_BID)", "Komplett" if submissions else "Ingen tilbud registrert"),
        ("Prosedyretype", "API (procedure)", "Komplett"),
        ("Kunngjøring", f"API {'(DOFFIN_NOTICE)' if doffin or publish else ''}", "Komplett" if doffin or publish else "Trenger bekreftelse"),
        ("Dialog/forhandlinger \u00a7 9-3", "**Manuelt**", "Ikke i API"),
        ("Avvisning formalfeil \u00a7 9-4", f"API ({len(rejections)} hendelser)" if rejections else "API (ingen hendelser)", "Trenger manuell klassifisering" if rejections else "Trenger bekreftelse"),
        ("Kvalifikasjonsvurdering", "**Manuelt**", "Ikke i API"),
        ("Avvisning leverandører \u00a7 9-5", f"API ({len(rejections)} hendelser)" if rejections else "API (ingen hendelser)", "Begrunnelse mangler" if rejections else "Trenger bekreftelse"),
        ("Avvisning tilbud \u00a7 9-6", "API (ingen hendelser)", "Trenger bekreftelse"),
        ("Ettersending/avklaring", f"API ({len(post_deadline)} meldinger etter frist)" if post_deadline else "API (ingen hendelser)", "Innhold mangler" if post_deadline else "Trenger bekreftelse"),
        ("Tilbud i vurdering", "API (SUBMIT_BID)", "Komplett" if submissions else "Ingen"),
        ("Valgte tilbud + begrunnelse", "**Manuelt**", "API har kun tenderIds"),
        ("Meddelelsesbrev/klagefrist", "**Manuelt**", "Kun flag, ikke datoer"),
        ("Underleverandører", "**Manuelt**", "Ikke i API"),
        ("Inhabilitet", "**Manuelt**", "Ikke i API"),
    ]
    _docx_add_table(doc, ["Seksjon", "Kilde", "Merknad"], rows)


# -- Main generator ----------------------------------------------------------

def generate_protokoll(procurement: dict, activities: list[dict]) -> str:
    """Generate a complete anskaffelsesprotokoll in markdown.

    Args:
        procurement: Full procurement object from Artifik API.
        activities: List of activity objects from get_procurement_activities.

    Returns:
        Markdown string with the protocol.
    """
    procedure = procurement.get("procedure") or ""
    seq_id = procurement.get("sequenceId") or procurement.get("name") or "Ukjent"
    today = datetime.now().strftime("%Y-%m-%d")

    sections = []

    sections.append(f"# ANSKAFFELSESPROTOKOLL — {seq_id}")
    sections.append("")
    sections.append(f"> **Generert fra API-data {today}.** Felter markert med `<!-- MANUELT -->` krever manuell utfylling.")
    sections.append("")

    sections.append(_section_general_info(procurement, activities))
    sections.append("---\n")
    sections.append(_section_procedure(procurement, activities))
    sections.append("---\n")
    sections.append(_section_formal_rejection(activities))
    sections.append("---\n")
    sections.append(_section_preliminary_qualification(procedure))
    sections.append("---\n")
    sections.append(_section_qualification())
    sections.append("---\n")
    sections.append(_section_supplier_rejection(activities))
    sections.append("---\n")
    sections.append(_section_supplier_selection(procedure, activities))
    sections.append("---\n")
    sections.append(_section_bid_rejection())
    sections.append("---\n")
    sections.append(_section_clarification(procurement, activities))
    sections.append("---\n")
    sections.append(_section_negotiations(procedure))
    sections.append("---\n")
    sections.append(_section_dialog(procedure))
    sections.append("---\n")
    sections.append(_section_bids_in_evaluation(activities))
    sections.append("---\n")
    sections.append(_section_award(procurement, activities))
    sections.append("---\n")
    sections.append(_section_framework_agreement(procurement))
    sections.append("---\n")
    sections.append(_section_other(procurement))
    sections.append("---\n")
    sections.append(_section_data_quality(procurement, activities))

    return "\n".join(sections) + "\n"


def generate_protokoll_docx(procurement: dict, activities: list[dict]) -> "DocxDocument":
    """Generate a complete anskaffelsesprotokoll as a Word document."""
    if not _HAS_DOCX:
        _die("python-docx er ikke installert. Kjør: pip install python-docx")

    procedure = procurement.get("procedure") or ""
    seq_id = procurement.get("sequenceId") or procurement.get("name") or "Ukjent"
    today = datetime.now().strftime("%d.%m.%Y")

    doc = DocxDocument()
    _docx_setup(doc)

    doc.add_heading(f"ANSKAFFELSESPROTOKOLL — {seq_id}", level=1)
    p = doc.add_paragraph()
    run = p.add_run(f"Generert fra API-data {today}. ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run2 = p.add_run("Felter med gul bakgrunn krever manuell utfylling.")
    run2.font.size = Pt(9)
    run2.font.italic = True
    run2.font.color.rgb = RGBColor(0x99, 0x66, 0x00)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "FFFF00")
    run2._element.get_or_add_rPr().append(shading)

    _docx_general_info(doc, procurement, activities)
    _docx_procedure(doc, procurement, activities)
    _docx_formal_rejection(doc, activities)
    _docx_preliminary_qualification(doc, procedure)
    _docx_qualification(doc)
    _docx_supplier_rejection(doc, activities)
    _docx_supplier_selection(doc, procedure, activities)
    _docx_bid_rejection(doc)
    _docx_clarification(doc, procurement, activities)
    _docx_negotiations(doc, procedure)
    _docx_dialog(doc, procedure)
    _docx_bids_in_evaluation(doc, activities)
    _docx_award(doc, procurement, activities)
    _docx_framework_agreement(doc, procurement)
    _docx_other(doc, procurement)
    _docx_data_quality(doc, procurement, activities)

    return doc


def generate_protokoll_docx_del2(procurement: dict, activities: list[dict]) -> "DocxDocument":
    """Generate anskaffelsesprotokoll for Del II (nasjonal, under EØS) as Word document."""
    if not _HAS_DOCX:
        _die("python-docx er ikke installert. Kjør: pip install python-docx")

    procedure = procurement.get("procedure") or ""
    seq_id = procurement.get("sequenceId") or procurement.get("name") or "Ukjent"
    today = datetime.now().strftime("%d.%m.%Y")

    doc = DocxDocument()
    _docx_setup(doc)

    doc.add_heading(f"ANSKAFFELSESPROTOKOLL for anskaffelser etter forskriften del II — {seq_id}", level=1)
    p = doc.add_paragraph()
    run = p.add_run(f"Generert fra API-data {today}. ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run2 = p.add_run("Felter med gul bakgrunn krever manuell utfylling.")
    run2.font.size = Pt(9)
    run2.font.italic = True
    run2.font.color.rgb = RGBColor(0x99, 0x66, 0x00)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "FFFF00")
    run2._element.get_or_add_rPr().append(shading)

    _docx_del2_general_info(doc, procurement, activities)
    _docx_del2_procedure(doc, procurement, activities)
    _docx_del2_dialog(doc, procurement, activities)
    _docx_del2_formal_rejection(doc, activities)
    _docx_del2_qualification(doc)
    _docx_del2_supplier_rejection(doc, activities)
    _docx_del2_supplier_selection(doc, procedure, activities)
    _docx_del2_bid_rejection(doc)
    _docx_del2_clarification(doc, procurement, activities)
    # Reuse _docx_bids_in_evaluation from Del III — same logic
    _docx_bids_in_evaluation(doc, activities)
    _docx_del2_award(doc, procurement, activities)
    _docx_del2_award_notification(doc, procurement)
    # Reuse _docx_framework_agreement from Del III — same logic
    _docx_framework_agreement(doc, procurement)
    _docx_del2_other(doc, procurement)
    _docx_del2_data_quality(doc, procurement, activities)

    return doc


# -- CLI ---------------------------------------------------------------------

def _print_summary(procurement: dict, activities: list[dict], path: str) -> None:
    """Print a summary of what was generated and what needs manual work."""
    submissions = _get_activities_by_action(activities, "SUBMIT_BID")
    rejections = _get_activities_by_action(activities, "REJECT_PARTICIPATION")
    doffin = _get_activities_by_action(activities, "DOFFIN_NOTICE_STATUS_PUBLISHED")
    publish = _get_activities_by_action(activities, "PUBLISH_TO_DOFFIN")
    awards = _get_activities_by_action(activities, "AWARDING_PARTICIPANTS")

    print(file=sys.stderr)
    print(f"  {_bold('Generert:')} {path}", file=sys.stderr)
    print(file=sys.stderr)

    # Auto-filled
    auto = []
    if procurement.get("about_procurer"):
        auto.append("Oppdragsgiver")
    if procurement.get("name"):
        auto.append("Beskrivelse")
    if procurement.get("procedure"):
        auto.append("Prosedyre")
    if doffin or publish:
        auto.append("Kunngjøring")
    if submissions:
        auto.append(f"{len(submissions)} tilbud mottatt")
    if rejections:
        auto.append(f"{len(rejections)} avvisninger")
    if awards:
        auto.append("Tildelingsdato")

    if auto:
        print(f"  {_green('Fylt fra API:')}  {', '.join(auto)}", file=sys.stderr)

    # Manual
    manual = [
        "Kvalifikasjonsvurdering",
        "Tildelingsbegrunnelse",
        "Delkontrakt-begrunnelse",
        "Meddelelsesbrev/karens",
    ]
    if rejections:
        manual.append("Avvisningsbegrunnelser")
    manual.extend(["Underleverandorer", "Inhabilitet"])

    print(
        f"  {_yellow('Trenger utfylling:')}  {', '.join(manual)}",
        file=sys.stderr,
    )

    # Count manual markers
    if path.endswith(".md"):
        result_text = Path(path).read_text()
        manual_count = result_text.count("<!-- MANUELT")
        if manual_count:
            print(
                f"\n  {_dim(f'Søk etter MANUELT i filen — {manual_count} steder trenger oppmerksomhet.')}",
                file=sys.stderr,
            )
    elif path.endswith(".docx"):
        print(
            f"\n  {_dim('Åpne filen — felter med gul bakgrunn trenger utfylling.')}",
            file=sys.stderr,
        )
    print(file=sys.stderr)


def main() -> None:
    import argparse

    parser = argparse.ArgumentParser(
        description="Generer anskaffelsesprotokoll fra Artifik API.",
    )
    parser.add_argument(
        "--id",
        type=int,
        help="Procurement ID (skipper interaktiv velging)",
    )
    parser.add_argument(
        "--list",
        action="store_true",
        dest="list_only",
        help="Bare list anskaffelser, ikke generer protokoll",
    )
    parser.add_argument(
        "-o", "--output",
        help="Output-fil (default: docs/protokoll-{sequenceId}.md)",
    )
    parser.add_argument(
        "--format",
        choices=["docx", "md"],
        default="docx",
        help="Output-format (default: docx)",
    )

    args = parser.parse_args()

    # Banner
    print(file=sys.stderr)
    print(f"  {_bold('Protokollgenerator')}", file=sys.stderr)
    print(f"  {_dim('Anskaffelsesprotokoll fra Artifik API')}", file=sys.stderr)
    print(file=sys.stderr)

    # Step 1: Connect
    _step(1, 3, "Kobler til Artifik API")
    client = _get_client()

    # Step 2: List
    _step(2, 3, "Henter anskaffelser")
    procurements = _list_procurements(client)

    if not procurements:
        _die("Ingen modne anskaffelser funnet (passert tilbudsfrist, ikke kansellert/mal).")

    if args.list_only:
        _print_procurement_list(procurements)
        return

    # Select procurement
    if args.id:
        matches = [p for p in procurements if p.get("id") == args.id]
        if not matches:
            all_procs = client.list_procurements()
            matches = [p for p in all_procs if p.get("id") == args.id]
            if not matches:
                _die(f"Fant ingen anskaffelse med ID {args.id}.")
        procurement = matches[0]
    else:
        procurement = _select_procurement(procurements)

    pid = procurement["id"]
    seq_id = procurement.get("sequenceId") or str(pid)
    name = procurement.get("name") or ""
    print(
        f"  {_green('>')} {_bold(seq_id)} — {name}",
        file=sys.stderr,
    )
    print(file=sys.stderr)

    # Step 3: Generate
    _step(3, 3, "Genererer protokoll")

    with _Spinner("Henter aktivitetslogg"):
        activities = client.get_procurement_activities(pid)
    _ok(f"{len(activities)} hendelser")

    fmt = args.format
    threshold = procurement.get("threshold") or ""
    is_del2 = threshold in ("below_eea_threshold_value", "national_threshold")

    if is_del2:
        _ok(f"Del II-protokoll (terskel: {THRESHOLD_SHORT.get(threshold, threshold)})")
    else:
        _ok(f"Del III-protokoll (terskel: {THRESHOLD_SHORT.get(threshold, threshold)})")

    if fmt == "md":
        result = generate_protokoll(procurement, activities)
        output_path = args.output or f"docs/protokoll-{seq_id.lower()}.md"
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        Path(output_path).write_text(result)
    else:
        if is_del2:
            doc = generate_protokoll_docx_del2(procurement, activities)
        else:
            doc = generate_protokoll_docx(procurement, activities)
        output_path = args.output or f"docs/protokoll-{seq_id.lower()}.docx"
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)

    _print_summary(procurement, activities, output_path)


if __name__ == "__main__":
    main()
